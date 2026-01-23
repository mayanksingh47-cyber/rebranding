"""
Reface - Brand transformation at the speed of scale.
DXC Enterprise Unified Document Rebranding Solution

Rebrands multiple documents at once using new DXC Dec 2025 templates:
- Word Documents (.docx)
- PDF Documents (.pdf)

Features:
- Upload multiple files at once (PDF and Word together)
- Batch processing with progress tracking
- Download all rebranded files as ZIP
- Individual file download options
- Processing summary and statistics

Author: Mayank Kumar
Cross Functional Capabilities - AI and Automation
Organization: DXC Technology
Version: 2.0.0 - Unified Platform
"""

import streamlit as st
import os
import io
import zipfile
import tempfile
import base64
import numpy as np
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import time

# ============================================================================
# PDF REBRANDING CLASS (Standalone - copied from dxc_pdf_rebranding.py)
# ============================================================================

# Check for PyMuPDF first (independent check - no poppler needed)
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Check for pdf2image (requires poppler)
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# Check for other dependencies
try:
    from PyPDF2 import PdfReader
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch, cm
    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import HexColor
    from PIL import Image, ImageDraw
    PDF_OTHER_DEPS_AVAILABLE = True
except ImportError:
    PDF_OTHER_DEPS_AVAILABLE = False

# Export availability flags
PDF_AVAILABLE = PYMUPDF_AVAILABLE or PDF2IMAGE_AVAILABLE
PDF_ALL_DEPS_AVAILABLE = PDF_AVAILABLE and PDF_OTHER_DEPS_AVAILABLE


class UnifiedPDFRebrander:
    """
    Rebrands PDF documents with DXC branding.
    
    CRITICAL: Uses SAME logic as Word rebranding:
    - REPLACES old DXC logos (black/purple) with new logo
    - REPLACES purple colors with orange/melon (#FF7E51)
    - REPLACES purple shapes/backgrounds with orange
    - REPLACES header with new DXC header
    - REPLACES footer with new DXC footer
    - All other content PRESERVED EXACTLY
    """
    
    def __init__(self, poppler_path: Optional[str] = None):
        # DXC Brand Colors (same as Word rebrander)
        self.DXC_DARK_NAVY = HexColor('#0E1020')
        self.DXC_ORANGE = HexColor('#EF8900')
        self.DXC_MELON = HexColor('#FF7E51')  # Melon color for replacing purple
        self.DXC_GRAY = HexColor('#6B7280')
        self.DXC_LIGHT_BG = HexColor('#F6F3F0')
        self.DXC_LIGHT_BLUE = HexColor('#87CEEB')  # Light blue (Sky Blue)
        self.DXC_DARK_BLUE = HexColor('#00008B')  # Dark blue
        
        # Poppler path (for Windows users who need to specify it - only if using pdf2image)
        self.poppler_path = poppler_path
        
        # Use PyMuPDF if available (preferred - no poppler needed)
        self.use_pymupdf = PYMUPDF_AVAILABLE
        
        # Page dimensions (A4)
        self.PAGE_WIDTH = A4[0]
        self.PAGE_HEIGHT = A4[1]
        
        # Header/Footer dimensions
        self.HEADER_HEIGHT = 1.0 * inch
        self.FOOTER_HEIGHT = 0.6 * inch
        self.LOGO_WIDTH = 4.5 * cm
        self.LOGO_HEIGHT = 1.0 * cm
        
        # Logo detection thresholds (same as Word rebrander)
        self.MAX_LOGO_WIDTH_CM = 10
        self.MAX_LOGO_HEIGHT_CM = 4
        
    def _find_logo(self) -> Optional[str]:
        """Find DXC logo file (same paths as Word rebrander)"""
        logo_paths = [
            'extracted_image1.png',
            'DXCLogo.png',
            'dxc_brand_logo.png',
            'AI Agent 1 - System Network and Services Agent/DXCLogo.png'
        ]
        
        for logo_path in logo_paths:
            if os.path.exists(logo_path):
                return logo_path
        return None
    
    def _is_purple_color(self, r: int, g: int, b: int) -> bool:
        """
        Detect purple colors (same logic as Word rebrander).
        Returns True if color is purple/violet/magenta.
        """
        r, g, b = int(r), int(g), int(b)
        r = max(0, min(255, r))
        g = max(0, min(255, g))
        b = max(0, min(255, b))
        
        is_classic_purple = (r > 80 and b > 80 and g < 100 and (r + b) > (2 * g + 50))
        is_violet = (r > 100 and b > 100 and abs(r - b) < 80 and g < max(r, b) * 0.6)
        is_magenta = (r > 100 and b > 100 and abs(r - b) < 80 and g < max(r, b) * 0.6)
        is_office_purple = (r >= 112 and r <= 142 and g >= 48 and g <= 78 and b >= 160 and b <= 190)
        
        return is_classic_purple or is_violet or is_magenta or is_office_purple
    
    def _replace_purple_colors_in_image(self, img: Image.Image, replacement_color: str = 'orange') -> Image.Image:
        """
        Replace purple colors with specified color in image.
        Uses FAST vectorized NumPy operations (no loops).
        
        Args:
            img: Input image
            replacement_color: 'orange' (default) or 'black' - color to replace purple with
        """
        img_array = np.array(img.convert('RGB'))
        r, g, b = img_array[:, :, 0], img_array[:, :, 1], img_array[:, :, 2]
        
        is_classic_purple = (r > 80) & (b > 80) & (g < 100) & ((r + b) > (2 * g + 50))
        is_violet = (r > 100) & (b > 100) & (np.abs(r - b) < 80) & (g < np.maximum(r, b) * 0.6)
        is_magenta = (r > 100) & (b > 100) & (np.abs(r - b) < 80) & (g < np.maximum(r, b) * 0.6)
        is_office_purple = (r >= 112) & (r <= 142) & (g >= 48) & (g <= 78) & (b >= 160) & (b <= 190)
        
        purple_mask = is_classic_purple | is_violet | is_magenta | is_office_purple
        
        # Replace ALL purple with specified color
        color_lower = replacement_color.lower()
        if color_lower == 'black':
            # Black = RGB 0, 0, 0
            img_array[purple_mask, 0] = 0   # R
            img_array[purple_mask, 1] = 0   # G
            img_array[purple_mask, 2] = 0   # B
        elif color_lower == 'light blue':
            # Light blue (Sky Blue) = RGB 135, 206, 235
            img_array[purple_mask, 0] = 135  # R
            img_array[purple_mask, 1] = 206  # G
            img_array[purple_mask, 2] = 235  # B
        elif color_lower == 'dark blue':
            # Dark blue = RGB 0, 0, 139
            img_array[purple_mask, 0] = 0    # R
            img_array[purple_mask, 1] = 0    # G
            img_array[purple_mask, 2] = 139  # B
        else:
            # Default: Orange/Melon (#FF7E51 = RGB 255, 126, 81)
            img_array[purple_mask, 0] = 255  # R
            img_array[purple_mask, 1] = 126  # G
            img_array[purple_mask, 2] = 81   # B
        
        return Image.fromarray(img_array, mode='RGB')
    
    def _detect_content_position(self, img: Image.Image) -> str:
        """Detect where content is positioned (left/right/center)."""
        width, height = img.size
        img_array = np.array(img.convert('RGB'))
        
        mid_x = width // 2
        left_half = img_array[:, :mid_x]
        right_half = img_array[:, mid_x:]
        
        left_non_white = np.sum((left_half[:, :, 0] < 240) | (left_half[:, :, 1] < 240) | (left_half[:, :, 2] < 240))
        right_non_white = np.sum((right_half[:, :, 0] < 240) | (right_half[:, :, 1] < 240) | (right_half[:, :, 2] < 240))
        
        total_pixels_left = left_half.shape[0] * left_half.shape[1]
        total_pixels_right = right_half.shape[0] * right_half.shape[1]
        
        left_ratio = left_non_white / total_pixels_left if total_pixels_left > 0 else 0
        right_ratio = right_non_white / total_pixels_right if total_pixels_right > 0 else 0
        
        if left_ratio > right_ratio * 1.5:
            return 'left'
        elif right_ratio > left_ratio * 1.5:
            return 'right'
        else:
            return 'center'
    
    def _remove_old_logos_from_header(self, img: Image.Image) -> Image.Image:
        """Remove old DXC logos from header, footer, and edges."""
        width, height = img.size
        img_array = np.array(img.convert('RGB'))
        
        estimated_dpi = height / (A4[1] / 72)
        if estimated_dpi < 100:
            estimated_dpi = 150
        
        max_logo_w = int(10 * estimated_dpi / 2.54)
        max_logo_h = int(4 * estimated_dpi / 2.54)
        min_logo_w = int(2 * estimated_dpi / 2.54)
        min_logo_h = int(0.5 * estimated_dpi / 2.54)
        
        regions_to_remove = []
        
        header_clear_height = int(1.5 * inch * estimated_dpi / 72)
        header_clear_width = int(18 * cm * estimated_dpi / 2.54)
        
        if header_clear_height > 0 and header_clear_width > 0:
            header_area = img_array[:header_clear_height, :min(header_clear_width, width)]
            if header_area.size > 0:
                r, g, b = header_area[:, :, 0], header_area[:, :, 1], header_area[:, :, 2]
                rgb_sum = r + g + b
                is_dark_text = (rgb_sum < 180)
                is_content = (rgb_sum < 600) & ~((r > 240) & (g > 240) & (b > 240))
                
                block_size = int(2 * cm * estimated_dpi / 2.54)
                for y in range(0, header_area.shape[0], block_size // 2):
                    for x in range(0, header_area.shape[1], block_size // 2):
                        x_end = min(x + block_size, header_area.shape[1])
                        y_end = min(y + block_size, header_area.shape[0])
                        block = header_area[y:y_end, x:x_end]
                        if block.size == 0:
                            continue
                        
                        r_b, g_b, b_b = block[:, :, 0], block[:, :, 1], block[:, :, 2]
                        rgb_sum_b = r_b + g_b + b_b
                        is_dark_b = (rgb_sum_b < 180)
                        is_content_b = (rgb_sum_b < 600) & ~((r_b > 240) & (g_b > 240) & (b_b > 240))
                        
                        dark_ratio = np.sum(is_dark_b) / block.shape[0] / block.shape[1] if block.shape[0] > 0 and block.shape[1] > 0 else 0
                        content_ratio = np.sum(is_content_b) / block.shape[0] / block.shape[1] if block.shape[0] > 0 and block.shape[1] > 0 else 0
                        
                        if dark_ratio > 0.01 or content_ratio > 0.10:
                            continue
                        else:
                            regions_to_remove.append((x, y, x_end - x, y_end - y))
        
        footer_clear_height = int(1.5 * inch * estimated_dpi / 72)
        footer_clear_width = int(18 * cm * estimated_dpi / 2.54)
        footer_start = height - footer_clear_height
        
        if footer_start < height and footer_start >= 0 and footer_clear_width > 0:
            footer_area = img_array[footer_start:, :min(footer_clear_width, width)]
            if footer_area.size > 0:
                r, g, b = footer_area[:, :, 0], footer_area[:, :, 1], footer_area[:, :, 2]
                rgb_sum = r + g + b
                is_orange = (r > 120) & (g > 30) & (b < 230) & (r > g * 0.3)
                orange_ratio = np.sum(is_orange) / footer_area.shape[0] / footer_area.shape[1] if footer_area.shape[0] > 0 and footer_area.shape[1] > 0 else 0
                
                if orange_ratio > 0.15:
                    pass
                else:
                    block_size = int(2 * cm * estimated_dpi / 2.54)
                    for y in range(0, footer_area.shape[0], block_size // 2):
                        for x in range(0, footer_area.shape[1], block_size // 2):
                            x_end = min(x + block_size, footer_area.shape[1])
                            y_end = min(y + block_size, footer_area.shape[0])
                            block = footer_area[y:y_end, x:x_end]
                            if block.size == 0:
                                continue
                            
                            r_b, g_b, b_b = block[:, :, 0], block[:, :, 1], block[:, :, 2]
                            rgb_sum_b = r_b + g_b + b_b
                            is_dark_b = (rgb_sum_b < 180)
                            is_content_b = (rgb_sum_b < 600) & ~((r_b > 240) & (g_b > 240) & (b_b > 240))
                            
                            dark_ratio = np.sum(is_dark_b) / block.shape[0] / block.shape[1] if block.shape[0] > 0 and block.shape[1] > 0 else 0
                            content_ratio = np.sum(is_content_b) / block.shape[0] / block.shape[1] if block.shape[0] > 0 and block.shape[1] > 0 else 0
                            
                            if dark_ratio > 0.01 or content_ratio > 0.10:
                                continue
                            else:
                                regions_to_remove.append((x, footer_start + y, x_end - x, y_end - y))
        
        result_img = img.copy()
        if result_img.mode != 'RGB':
            result_img = result_img.convert('RGB')
        
        img_array = np.array(result_img)
        for x, y, w, h in regions_to_remove:
            x1, y1 = max(0, int(x)), max(0, int(y))
            x2, y2 = min(width, int(x + w)), min(height, int(y + h))
            if x2 > x1 and y2 > y1:
                img_array[y1:y2, x1:x2] = [255, 255, 255]
        
        result_img = Image.fromarray(img_array)
        draw = ImageDraw.Draw(result_img)
        for x, y, w, h in regions_to_remove:
            x1, y1 = max(0, int(x)), max(0, int(y))
            x2, y2 = min(width, int(x + w)), min(height, int(y + h))
            if x2 > x1 and y2 > y1:
                draw.rectangle([x1, y1, x2, y2], fill=(255, 255, 255), outline=(255, 255, 255), width=0)
        
        return result_img
    
    def _add_header_footer_to_image(self, img: Image.Image, logo_path: Optional[str]) -> Image.Image:
        """Add DXC header logo to processed page image."""
        width, height = img.size
        
        result_img = self._remove_old_logos_from_header(img)
        content_pos = self._detect_content_position(result_img)
        
        if logo_path and os.path.exists(logo_path):
            try:
                logo_img = Image.open(logo_path)
                if logo_img.mode != 'RGBA':
                    logo_img = logo_img.convert('RGBA')
                
                estimated_dpi = height / (A4[1] / 72)
                if estimated_dpi < 100:
                    estimated_dpi = 200
                
                logo_width_cm = 4.5
                logo_width_px = int(logo_width_cm * estimated_dpi / 2.54)
                
                original_width, original_height = logo_img.size
                aspect_ratio = original_height / original_width
                logo_height_px = int(logo_width_px * aspect_ratio)
                
                max_height_px = int(1.5 * cm * estimated_dpi / 2.54)
                if logo_height_px > max_height_px:
                    logo_height_px = max_height_px
                    logo_width_px = int(logo_height_px / aspect_ratio)
                
                logo_img = logo_img.resize((logo_width_px, logo_height_px), Image.Resampling.LANCZOS)
                
                header_zone_height = int(1.0 * inch * estimated_dpi / 72)
                logo_y = int(0.2 * inch * estimated_dpi / 72)
                
                margin_px = int(0.5 * inch * estimated_dpi / 72)
                if content_pos == 'left':
                    logo_x = width - logo_width_px - margin_px
                elif content_pos == 'right':
                    logo_x = margin_px
                else:
                    logo_x = margin_px
                
                logo_x = max(0, min(logo_x, width - logo_width_px))
                logo_y = max(0, min(logo_y, header_zone_height - logo_height_px))
                
                if logo_y + logo_height_px > header_zone_height:
                    logo_y = header_zone_height - logo_height_px - int(0.1 * inch * estimated_dpi / 72)
                    if logo_y < 0:
                        logo_y = int(0.1 * inch * estimated_dpi / 72)
                
                if logo_img.mode == 'RGBA':
                    logo_array = np.array(logo_img, dtype=np.uint8)
                    r, g, b, a = logo_array[:, :, 0], logo_array[:, :, 1], logo_array[:, :, 2], logo_array[:, :, 3]
                    is_black = (r < 50) & (g < 50) & (b < 50)
                    a[is_black] = 0
                    logo_array[:, :, 3] = a
                    logo_img_clean = Image.fromarray(logo_array, mode='RGBA')
                    result_img.paste(logo_img_clean, (logo_x, logo_y), logo_img_clean.split()[3])
                else:
                    logo_rgba = logo_img.convert('RGBA')
                    logo_array = np.array(logo_rgba, dtype=np.uint8)
                    r, g, b, a = logo_array[:, :, 0], logo_array[:, :, 1], logo_array[:, :, 2], logo_array[:, :, 3]
                    is_black = (r < 50) & (g < 50) & (b < 50)
                    a[is_black] = 0
                    logo_array[:, :, 3] = a
                    logo_rgba_clean = Image.fromarray(logo_array, mode='RGBA')
                    result_img.paste(logo_rgba_clean, (logo_x, logo_y), logo_rgba_clean.split()[3])
            except Exception as e:
                import sys
                sys.stderr.write(f"Warning: Could not add logo: {str(e)}\n")
        
        return result_img
    
    def rebrand_pdf(self, source_file: bytes, source_filename: str, replacement_color: str = 'orange') -> Tuple[bytes, str]:
        """
        Rebrand a PDF document with DXC branding.
        
        Args:
            source_file: PDF file bytes
            source_filename: Original filename
            replacement_color: 'orange' (default) or 'black' - color to replace purple shapes with
        """
        if not PYMUPDF_AVAILABLE and not PDF2IMAGE_AVAILABLE:
            raise RuntimeError(
                "No PDF conversion library available. Please install one:\n"
                "pip install PyMuPDF (recommended - no poppler needed)\n"
                "OR: pip install pdf2image (requires poppler)"
            )
        
        if not PDF_OTHER_DEPS_AVAILABLE:
            raise RuntimeError(
                "Missing required dependencies. Please install:\n"
                "pip install PyPDF2 reportlab Pillow numpy"
            )
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            tmp.write(source_file)
            source_path = tmp.name
        
        try:
            logo_path = self._find_logo()
            
            page_images = []
            if PYMUPDF_AVAILABLE:
                doc = fitz.open(source_path)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    pix = page.get_pixmap(matrix=fitz.Matrix(150/72, 150/72))
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    page_images.append(img)
                doc.close()
            elif PDF2IMAGE_AVAILABLE:
                images = convert_from_path(
                    source_path,
                    dpi=150,
                    poppler_path=self.poppler_path
                )
                page_images = images
            
            processed_images = []
            for img in page_images:
                img = self._replace_purple_colors_in_image(img, replacement_color)
                img = self._add_header_footer_to_image(img, logo_path)
                processed_images.append(img)
            
            output_buffer = io.BytesIO()
            c = canvas.Canvas(output_buffer, pagesize=A4)
            
            for img in processed_images:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                    img.save(tmp_img.name, format='PNG')
                    tmp_img_path = tmp_img.name
                
                try:
                    c.drawImage(tmp_img_path, 0, 0, width=A4[0], height=A4[1], preserveAspectRatio=True)
                    c.showPage()
                finally:
                    if os.path.exists(tmp_img_path):
                        os.unlink(tmp_img_path)
            
            c.save()
            output_buffer.seek(0)
            
            output_name = f"DXC_Rebranded_{Path(source_filename).stem}_{datetime.now().strftime('%Y%m%d')}.pdf"
            
            return output_buffer.getvalue(), output_name
            
        finally:
            if os.path.exists(source_path):
                os.unlink(source_path)


# ============================================================================
# WORD REBRANDING CLASS (Standalone - copied from dxc_document_rebranding.py)
# ============================================================================

# Optional imports for Word
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class UnifiedWordRebrander:
    """
    Rebrands Word documents with DXC branding.
    
    CRITICAL: Only changes HEADERS and FOOTERS.
    Document content, fonts, tables, and formatting are PRESERVED EXACTLY as-is.
    """
    
    def __init__(self):
        self.template_path = 'DXC Word_A4_DEC 25.dotx'
        self.template_folder = 'DXC Word_DEC 25'
        
        # DXC Brand Colors
        self.DXC_DARK_NAVY = RGBColor(0x0E, 0x10, 0x20)
        self.DXC_ORANGE = RGBColor(0xEF, 0x89, 0x00)
        self.DXC_LIGHT_BG = RGBColor(0xF6, 0xF3, 0xF0)
        self.DXC_GRAY = RGBColor(0x6B, 0x72, 0x80)
        self.DXC_MELON = RGBColor(0xFF, 0x7E, 0x51)  # Melon color for replacing purple
        self.DXC_BLACK = RGBColor(0x00, 0x00, 0x00)  # Black color for replacing purple
        self.DXC_LIGHT_BLUE = RGBColor(0x87, 0xCE, 0xEB)  # Light blue (Sky Blue)
        self.DXC_DARK_BLUE = RGBColor(0x00, 0x00, 0x8B)  # Dark blue
    
    def _get_replacement_color(self, replacement_color: str = 'orange'):
        """Get the replacement color RGBColor object based on the color choice."""
        color_lower = replacement_color.lower()
        if color_lower == 'black':
            return self.DXC_BLACK
        elif color_lower == 'light blue':
            return self.DXC_LIGHT_BLUE
        elif color_lower == 'dark blue':
            return self.DXC_DARK_BLUE
        else:
            return self.DXC_MELON
    
    def _get_replacement_color_hex(self, replacement_color: str = 'orange'):
        """Get the replacement color hex string based on the color choice."""
        color_lower = replacement_color.lower()
        if color_lower == 'black':
            return '000000'
        elif color_lower == 'light blue':
            return '87CEEB'
        elif color_lower == 'dark blue':
            return '00008B'
        else:
            return 'FF7E51'
    
    def _is_dark_color(self, color_hex: str) -> bool:
        """
        Determine if a color is dark or light.
        Uses relative luminance formula to determine brightness.
        
        Args:
            color_hex: Hex color string (with or without #)
            
        Returns:
            True if color is dark, False if light
        """
        try:
            # Clean hex string
            hex_str = color_hex.lstrip('#').upper()
            if len(hex_str) < 6:
                return False
            
            # Convert to RGB
            r = int(hex_str[0:2], 16)
            g = int(hex_str[2:4], 16)
            b = int(hex_str[4:6], 16)
            
            # Calculate relative luminance (perceived brightness)
            # Using standard formula: 0.299*R + 0.587*G + 0.114*B
            luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255.0
            
            # If luminance is less than 0.5, consider it dark
            return luminance < 0.5
        except:
            return False
    
    def _get_text_color_for_background(self, background_color: str) -> RGBColor:
        """
        Get appropriate text color (white or black) based on background color.
        
        Args:
            background_color: Background color hex string
            
        Returns:
            RGBColor object (white for dark backgrounds, black for light backgrounds)
        """
        if self._is_dark_color(background_color):
            return RGBColor(0xFF, 0xFF, 0xFF)  # White text for dark backgrounds
        else:
            return RGBColor(0x00, 0x00, 0x00)  # Black text for light backgrounds
    
    def _find_logo(self) -> Optional[str]:
        """Find DXC logo file"""
        logo_paths = [
            'extracted_image1.png',
            'DXCLogo.png',
            'dxc_brand_logo.png',
            'AI Agent 1 - System Network and Services Agent/DXCLogo.png'
        ]
        
        for logo_path in logo_paths:
            if os.path.exists(logo_path):
                return logo_path
        return None
    
    def _is_purple_color(self, color):
        """
        Check if a color is purple/violet (various shades).
        Covers: standard purple, violet, magenta, Office purple themes.
        """
        if color is None:
            return False
        try:
            color_str = str(color).upper()
            
            # Known purple/violet hex codes used in Office documents
            # Including DXC old brand purple colors
            known_purples = [
                '7030A0',  # Standard Office purple
                '800080',  # Purple
                '660066',  # Dark purple
                '9933FF',  # Bright purple
                '663399',  # Rebecca purple
                '8B008B',  # Dark magenta
                '9400D3',  # Dark violet
                'BA55D3',  # Medium orchid
                '9932CC',  # Dark orchid
                '8A2BE2',  # Blue violet
                'A020F0',  # Purple
                '6A0DAD',  # Violet
                '551A8B',  # Purple4
                '9370DB',  # Medium purple
                '7B68EE',  # Medium slate blue
                'EE82EE',  # Violet
                'DA70D6',  # Orchid
                'FF00FF',  # Magenta/Fuchsia
                'C71585',  # Medium violet red
                '5F249F',  # DXC old brand purple (common in documents)
                '8064A2',  # Light purple
                '60497A',  # Dark purple variant
            ]
            
            # Check if it's a known purple
            if color_str in known_purples:
                return True
            
            # Parse RGB values
            r = int(color_str[0:2], 16)
            g = int(color_str[2:4], 16)
            b = int(color_str[4:6], 16)
            
            # Broader purple detection (more aggressive):
            # 1. Classic purple: high R and B, low G
            # 2. Violet shades: B > R, both higher than G
            # 3. Magenta shades: R and B similar, both much higher than G
            # 4. DXC purple (5F249F): R=95, G=36, B=159
            
            is_classic_purple = (r > 80 and b > 80 and g < 100 and (r + b) > (2 * g + 50))
            is_violet = (b > r and b > 100 and g < 80 and r > 50)
            is_magenta = (r > 100 and b > 100 and abs(r - b) < 80 and g < max(r, b) * 0.6)
            is_office_purple = (r > 60 and b > 120 and g < 80)  # Catches 7030A0 type
            # DXC purple detection: R around 90-100, G around 30-50, B around 150-170
            is_dxc_purple = (r >= 85 and r <= 105 and g >= 25 and g <= 55 and b >= 140 and b <= 180)
            
            return is_classic_purple or is_violet or is_magenta or is_office_purple or is_dxc_purple
        except:
            return False
    
    def _replace_run_purple_color(self, run, replacement_color: str = 'orange'):
        """Helper to replace purple color in a single run."""
        replacement_rgb = self._get_replacement_color(replacement_color)
        replacement_hex = self._get_replacement_color_hex(replacement_color)
        
        try:
            # Check direct RGB color
            if run.font.color and run.font.color.rgb:
                color_hex = str(run.font.color.rgb)
                if self._is_purple_color(color_hex):
                    run.font.color.rgb = replacement_rgb
                    return True
            
            # Also check theme color - purple themes should be replaced
            if run.font.color and run.font.color.theme_color:
                theme = str(run.font.color.theme_color)
                # Theme colors that might be purple: ACCENT_4, ACCENT_5, etc.
                if 'ACCENT' in theme or 'PURPLE' in theme.upper():
                    run.font.color.rgb = replacement_rgb
                    return True
        except:
            pass
        
        # Try to access color via XML element directly (AGGRESSIVE)
        try:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                color_elem = rPr.find(qn('w:color'))
                if color_elem is not None:
                    val = color_elem.get(qn('w:val'))
                    # Also check 'val' attribute directly (without namespace)
                    if not val:
                        val = color_elem.get('val')
                    if val:
                        # Remove # if present
                        clean_val = val.lstrip('#').upper()
                        if self._is_purple_color(clean_val):
                            run.font.color.rgb = replacement_rgb
                            # Also update XML directly
                            color_elem.set(qn('w:val'), replacement_hex)
                            return True
        except:
            pass
        
        # Additional XML scan - check all color elements in run
        try:
            for color_elem in run._element.iter(qn('w:color')):
                val = color_elem.get(qn('w:val')) or color_elem.get('val')
                if val:
                    clean_val = val.lstrip('#').upper()
                    if self._is_purple_color(clean_val):
                        run.font.color.rgb = replacement_rgb
                        color_elem.set(qn('w:val'), replacement_hex)
                        return True
        except:
            pass
        
        return False
    
    def _replace_purple_in_numbering(self, doc: Document, replacement_color: str = 'orange'):
        """
        Replace purple colors in ALL numbering/list definitions.
        Handles all font types and all numbering levels.
        
        Args:
            doc: Word document
            replacement_color: 'orange' (default) or 'black' - color to replace purple with
        """
        replacement_hex = self._get_replacement_color_hex(replacement_color)
        try:
            # Access the numbering part of the document
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return
            
            numbering_xml = numbering_part._element
            
            # Namespace for Word ML
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            # Replace ALL color elements anywhere in numbering XML
            for color_elem in numbering_xml.iter('{%s}color' % w_ns):
                val = color_elem.get('{%s}val' % w_ns)
                if val and self._is_purple_color(val):
                    color_elem.set('{%s}val' % w_ns, replacement_hex)
                # Also check without namespace prefix
                val2 = color_elem.get('val')
                if val2 and self._is_purple_color(val2):
                    color_elem.set('val', replacement_hex)
            
            # Process all abstractNum definitions (base numbering styles)
            for abstractNum in numbering_xml.iter('{%s}abstractNum' % w_ns):
                for lvl in abstractNum.iter('{%s}lvl' % w_ns):
                    # Check rPr in each level
                    for rPr in lvl.iter('{%s}rPr' % w_ns):
                        for color_elem in rPr.iter('{%s}color' % w_ns):
                            val = color_elem.get('{%s}val' % w_ns) or color_elem.get('val')
                            if val and self._is_purple_color(val):
                                color_elem.set('{%s}val' % w_ns, replacement_hex)
            
            # Process all num definitions (actual numbering instances)
            for num in numbering_xml.iter('{%s}num' % w_ns):
                for lvlOverride in num.iter('{%s}lvlOverride' % w_ns):
                    for lvl in lvlOverride.iter('{%s}lvl' % w_ns):
                        for rPr in lvl.iter('{%s}rPr' % w_ns):
                            for color_elem in rPr.iter('{%s}color' % w_ns):
                                val = color_elem.get('{%s}val' % w_ns) or color_elem.get('val')
                                if val and self._is_purple_color(val):
                                    color_elem.set('{%s}val' % w_ns, replacement_hex)
                                    
        except Exception as e:
            pass  # Numbering part might not exist
    
    def _replace_purple_in_styles(self, doc: Document, replacement_color: str = 'orange'):
        """
        Replace purple colors in ALL document styles (headings, TOC, lists, etc.).
        Covers all font types and style definitions.
        
        Args:
            doc: Word document
            replacement_color: 'orange' (default) or 'black' - color to replace purple with
        """
        replacement_rgb = self._get_replacement_color(replacement_color)
        replacement_hex = self._get_replacement_color_hex(replacement_color)
        try:
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            # Process all styles via python-docx API
            for style in doc.styles:
                try:
                    if hasattr(style, 'font') and style.font and style.font.color:
                        if style.font.color.rgb:
                            color_hex = str(style.font.color.rgb)
                            if self._is_purple_color(color_hex):
                                style.font.color.rgb = replacement_rgb
                except:
                    pass
            
            # Direct XML modification for ALL color elements in styles
            styles_element = doc.styles.element
            
            # Replace all color elements
            for color_elem in styles_element.iter('{%s}color' % w_ns):
                val = color_elem.get('{%s}val' % w_ns) or color_elem.get('val')
                if val and self._is_purple_color(val):
                    color_elem.set('{%s}val' % w_ns, replacement_hex)
            
            # Process each style element individually
            purple_style_names = []
            for style_elem in styles_element.iter('{%s}style' % w_ns):
                style_name_attr = style_elem.get('{%s}styleId' % w_ns) or style_elem.get('styleId')
                is_purple_style = False
                
                # Check rPr (run properties)
                for rPr in style_elem.iter('{%s}rPr' % w_ns):
                    for color_elem in rPr.iter('{%s}color' % w_ns):
                        val = color_elem.get('{%s}val' % w_ns) or color_elem.get('val')
                        if val and self._is_purple_color(val):
                            color_elem.set('{%s}val' % w_ns, replacement_hex)
                            is_purple_style = True
                
                # Check pPr (paragraph properties)
                for pPr in style_elem.iter('{%s}pPr' % w_ns):
                    for rPr in pPr.iter('{%s}rPr' % w_ns):
                        for color_elem in rPr.iter('{%s}color' % w_ns):
                            val = color_elem.get('{%s}val' % w_ns) or color_elem.get('val')
                            if val and self._is_purple_color(val):
                                color_elem.set('{%s}val' % w_ns, replacement_hex)
                                is_purple_style = True
                
                # Track purple styles to apply to all paragraphs using them
                if is_purple_style and style_name_attr:
                    purple_style_names.append(style_name_attr)
                    
        except Exception as e:
            pass
        
        # CRITICAL: After updating styles, force-apply orange to all paragraphs using purple styles
        # This ensures ALL runs in those paragraphs get the orange color (fixes partial conversion issue)
        try:
            for style_name in purple_style_names:
                try:
                    for para in doc.paragraphs:
                        if para.style.name == style_name:
                            # This paragraph uses the purple style - apply replacement color to ALL runs
                            for run in para.runs:
                                try:
                                    run.font.color.rgb = replacement_rgb
                                except:
                                    pass
                except:
                    pass
        except:
            pass
    
    def _replace_purple_in_document_xml(self, doc: Document):
        """
        Direct XML scan to replace ALL purple colors in the entire document.
        This is a comprehensive fallback to catch any missed purple text.
        CRITICAL: Processes ALL paragraphs, headings, tables, headers, footers.
        """
        try:
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            
            # Get the main document element
            body = doc.element.body
            
            # Replace ALL color elements in the document body (text color)
            for color_elem in body.iter('{%s}color' % w_ns):
                val = color_elem.get('{%s}val' % w_ns) or color_elem.get('val')
                if val:
                    clean_val = val.lstrip('#').upper()
                    if self._is_purple_color(clean_val):
                        color_elem.set('{%s}val' % w_ns, 'FF7E51')
                        if color_elem.get('val'):
                            color_elem.set('val', 'FF7E51')
            
            # Also process ALL sections (headers, footers) for purple colors
            for section in doc.sections:
                # Main header
                try:
                    for color_elem in section.header._element.iter('{%s}color' % w_ns):
                        val = color_elem.get('{%s}val' % w_ns) or color_elem.get('val')
                        if val:
                            clean_val = val.lstrip('#').upper()
                            if self._is_purple_color(clean_val):
                                color_elem.set('{%s}val' % w_ns, 'FF7E51')
                                if color_elem.get('val'):
                                    color_elem.set('val', 'FF7E51')
                except:
                    pass
                
                # First page header
                try:
                    for color_elem in section.first_page_header._element.iter('{%s}color' % w_ns):
                        val = color_elem.get('{%s}val' % w_ns) or color_elem.get('val')
                        if val:
                            clean_val = val.lstrip('#').upper()
                            if self._is_purple_color(clean_val):
                                color_elem.set('{%s}val' % w_ns, 'FF7E51')
                                if color_elem.get('val'):
                                    color_elem.set('val', 'FF7E51')
                except:
                    pass
                        
        except Exception as e:
            pass
    
    def _replace_purple_backgrounds(self, doc: Document, replacement_color: str = 'orange'):
        """
        Replace purple background/shading colors with specified color.
        Handles table cell backgrounds, paragraph shading, etc.
        
        Args:
            doc: Word document
            replacement_color: 'orange' (default) or 'black' - color to replace purple with
        """
        replacement_hex = self._get_replacement_color_hex(replacement_color)
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        
        # Process table cell backgrounds
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    try:
                        # Access cell shading via XML
                        tc = cell._tc
                        tcPr = tc.find('{%s}tcPr' % w_ns)
                        background_replaced = False
                        if tcPr is not None:
                            shd = tcPr.find('{%s}shd' % w_ns)
                            if shd is not None:
                                # Check fill color
                                fill = shd.get('{%s}fill' % w_ns) or shd.get('fill')
                                if fill and self._is_purple_color(fill):
                                    shd.set('{%s}fill' % w_ns, replacement_hex)
                                    background_replaced = True
                                # Check color attribute
                                color = shd.get('{%s}color' % w_ns) or shd.get('color')
                                if color and self._is_purple_color(color):
                                    shd.set('{%s}color' % w_ns, replacement_hex)
                                    background_replaced = True
                        
                        # Adapt text color based on new background
                        if background_replaced:
                            text_color = self._get_text_color_for_background(replacement_hex)
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    if run.font.color is None or not run.font.color.rgb:
                                        run.font.color.rgb = text_color
                                    # Also update if text color is purple (being replaced)
                                    elif run.font.color.rgb:
                                        color_hex = str(run.font.color.rgb)
                                        clean_hex = color_hex.lstrip('#').upper()
                                        if len(clean_hex) >= 6 and self._is_purple_color(clean_hex[:6]):
                                            run.font.color.rgb = text_color
                    except:
                        pass
        
        # Process paragraph shading
        for para in doc.paragraphs:
            try:
                pPr = para._element.find('{%s}pPr' % w_ns)
                background_replaced = False
                if pPr is not None:
                    shd = pPr.find('{%s}shd' % w_ns)
                    if shd is not None:
                        fill = shd.get('{%s}fill' % w_ns) or shd.get('fill')
                        if fill and self._is_purple_color(fill):
                            shd.set('{%s}fill' % w_ns, replacement_hex)
                            background_replaced = True
                        color = shd.get('{%s}color' % w_ns) or shd.get('color')
                        if color and self._is_purple_color(color):
                            shd.set('{%s}color' % w_ns, replacement_hex)
                            background_replaced = True
                
                # Adapt text color based on new background
                if background_replaced:
                    text_color = self._get_text_color_for_background(replacement_hex)
                    for run in para.runs:
                        if run.font.color is None or not run.font.color.rgb:
                            run.font.color.rgb = text_color
                        # Also update if text color is purple (being replaced)
                        elif run.font.color.rgb:
                            color_hex = str(run.font.color.rgb)
                            clean_hex = color_hex.lstrip('#').upper()
                            if len(clean_hex) >= 6 and self._is_purple_color(clean_hex[:6]):
                                run.font.color.rgb = text_color
            except:
                pass
        
        # Scan entire document for any shading elements
        try:
            body = doc.element.body
            for shd in body.iter('{%s}shd' % w_ns):
                try:
                    background_replaced = False
                    # Replace fill color if purple
                    fill = shd.get('{%s}fill' % w_ns) or shd.get('fill')
                    if fill and self._is_purple_color(fill):
                        shd.set('{%s}fill' % w_ns, replacement_hex)
                        if shd.get('fill'):
                            shd.set('fill', replacement_hex)
                        background_replaced = True
                    
                    # Replace color if purple
                    color = shd.get('{%s}color' % w_ns) or shd.get('color')
                    if color and self._is_purple_color(color):
                        shd.set('{%s}color' % w_ns, replacement_hex)
                        if shd.get('color'):
                            shd.set('color', replacement_hex)
                        background_replaced = True
                            
                    # Also check themeFill and themeColor
                    for attr in ['themeFill', 'themeColor']:
                        val = shd.get('{%s}%s' % (w_ns, attr)) or shd.get(attr)
                        if val and 'accent' in val.lower():
                            # Remove theme reference and set direct color
                            shd.set('{%s}fill' % w_ns, replacement_hex)
                            background_replaced = True
                except:
                    pass
        except:
            pass
        
        # Also check styles for background colors
        try:
            styles_element = doc.styles.element
            for shd in styles_element.iter('{%s}shd' % w_ns):
                try:
                    fill = shd.get('{%s}fill' % w_ns) or shd.get('fill')
                    if fill and self._is_purple_color(fill):
                        shd.set('{%s}fill' % w_ns, replacement_hex)
                except:
                    pass
        except:
            pass
        
        # Replace purple shape/textbox background fills
        self._replace_purple_shape_fills(doc, replacement_color)
    
    def _replace_purple_shape_fills(self, doc: Document, replacement_color: str = 'orange'):
        """
        Replace purple background fills in shapes, textboxes, and drawing elements.
        Aggressively searches entire document XML for ANY purple color references.
        Handles both direct RGB colors AND theme color references.
        """
        replacement_hex = self._get_replacement_color_hex(replacement_color)
        
        # DrawingML namespace
        a_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
        
        def replace_in_element_tree(root_element):
            """Recursively search and replace purple colors in element tree."""
            if root_element is None:
                return
                
            # Search ALL elements in the tree
            for elem in list(root_element.iter()):
                try:
                    tag_name = elem.tag.split('}')[-1].lower() if '}' in elem.tag else elem.tag.lower()
                    
                    # HANDLE THEME COLORS: Replace schemeClr with direct srgbClr
                    if tag_name == 'schemeclr':
                        val = elem.get('val', '')
                        # Purple theme colors: accent4, accent5, accent6, dk2, etc.
                        if any(x in val.lower() for x in ['accent', 'dk1', 'dk2', 'tx1', 'tx2']):
                            parent = elem.getparent()
                            if parent is not None:
                                # Replace schemeClr with srgbClr replacement color
                                parent.remove(elem)
                                new_srgb = OxmlElement('a:srgbClr')
                                new_srgb.set('val', replacement_hex)
                                parent.append(new_srgb)
                    
                    # Check element tag for color-related names
                    if any(x in tag_name for x in ['srgbclr', 'solidfill', 'fill', 'bgcolor', 'color']):
                        # Check 'val' attribute (DrawingML colors)
                        val = elem.get('val')
                        if val:
                            # Remove # if present and check
                            clean_val = str(val).lstrip('#').upper()
                            if len(clean_val) >= 6:
                                clean_val = clean_val[:6]  # Take first 6 chars
                                if self._is_purple_color(clean_val):
                                    elem.set('val', replacement_hex)
                    
                    # Check ALL attributes of EVERY element for purple colors
                    for attr_name, attr_val in list(elem.attrib.items()):
                        if attr_val:
                            # Skip position-related attributes (preserve shape positions)
                            if attr_name.lower() in ['cx', 'cy', 'x', 'y', 'left', 'top', 'width', 'height']:
                                continue
                            
                            # Remove # prefix if present
                            clean_val = str(attr_val).lstrip('#').upper()
                            # Check if it looks like a hex color (6 chars, alphanumeric)
                            if len(clean_val) >= 6:
                                clean_val = clean_val[:6]  # Take first 6 chars
                                if clean_val.isalnum() and self._is_purple_color(clean_val):
                                    # Preserve # if original had it
                                    if str(attr_val).startswith('#'):
                                        elem.set(attr_name, '#' + replacement_hex)
                                    else:
                                        elem.set(attr_name, replacement_hex)
                except:
                    pass
        
        def replace_theme_colors_aggressive(root_element):
            """Force replace ALL solidFill elements that use theme colors or purple colors."""
            if root_element is None:
                return
            
            a_ns_full = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            
            # Find ALL solidFill elements
            for solid_fill in list(root_element.iter('{%s}solidFill' % a_ns_full)):
                try:
                    # Check if it has a schemeClr child (theme color)
                    scheme_clr = solid_fill.find('{%s}schemeClr' % a_ns_full)
                    if scheme_clr is not None:
                        val = scheme_clr.get('val', '')
                        # Replace ALL theme colors (accent colors often are purple)
                        # OR if it's specifically a purple theme
                        if any(x in val.lower() for x in ['accent', 'dk1', 'dk2', 'tx1', 'tx2']) or 'purple' in val.lower():
                            # This is a theme color - replace entirely with replacement color
                            solid_fill.remove(scheme_clr)
                            new_srgb = OxmlElement('a:srgbClr')
                            new_srgb.set('val', replacement_hex)
                            solid_fill.append(new_srgb)
                    
                    # Also check for srgbClr and replace if purple
                    srgb_clr = solid_fill.find('{%s}srgbClr' % a_ns_full)
                    if srgb_clr is not None:
                        val = srgb_clr.get('val')
                        if val and self._is_purple_color(val):
                            srgb_clr.set('val', replacement_hex)
                except:
                    pass
            
            # Find ALL spPr (shape properties) and replace fills if purple
            for spPr in list(root_element.iter('{%s}spPr' % a_ns_full)):
                try:
                    solid_fill = spPr.find('{%s}solidFill' % a_ns_full)
                    if solid_fill is not None:
                        # Check if current fill is purple
                        scheme_clr = solid_fill.find('{%s}schemeClr' % a_ns_full)
                        srgb_clr = solid_fill.find('{%s}srgbClr' % a_ns_full)
                        prst_clr = solid_fill.find('{%s}prstClr' % a_ns_full)  # Preset colors
                        
                        is_purple = False
                        if scheme_clr is not None:
                            val = scheme_clr.get('val', '')
                            # Replace ALL accent colors and theme colors (often purple)
                            if any(x in val.lower() for x in ['accent', 'dk1', 'dk2', 'tx1', 'tx2']) or 'purple' in val.lower():
                                is_purple = True
                        elif srgb_clr is not None:
                            val = srgb_clr.get('val')
                            if val:
                                clean_val = str(val).lstrip('#').upper()
                                if len(clean_val) >= 6:
                                    clean_val = clean_val[:6]
                                    if self._is_purple_color(clean_val):
                                        is_purple = True
                        elif prst_clr is not None:
                            # Preset colors might be purple - check val attribute
                            val = prst_clr.get('val', '')
                            if 'purple' in val.lower() or 'violet' in val.lower():
                                is_purple = True
                        
                        # If purple, replace with replacement color
                        if is_purple:
                            # Clear and recreate with replacement color
                            for child in list(solid_fill):
                                solid_fill.remove(child)
                            new_srgb = OxmlElement('a:srgbClr')
                            new_srgb.set('val', replacement_hex)
                            solid_fill.append(new_srgb)
                except:
                    pass
            
            # ALSO check for gradient fills that might have purple stops
            for gradFill in list(root_element.iter('{%s}gradFill' % a_ns_full)):
                try:
                    # Check gradient stops for purple colors
                    gs_list = gradFill.find('{%s}gsLst' % a_ns_full)
                    if gs_list is not None:
                        for gs in gs_list.iter('{%s}gs' % a_ns_full):
                            solid_fill = gs.find('{%s}solidFill' % a_ns_full)
                            if solid_fill is not None:
                                scheme_clr = solid_fill.find('{%s}schemeClr' % a_ns_full)
                                srgb_clr = solid_fill.find('{%s}srgbClr' % a_ns_full)
                                is_purple = False
                                if scheme_clr is not None:
                                    val = scheme_clr.get('val', '')
                                    if any(x in val.lower() for x in ['accent', 'dk1', 'dk2', 'tx1', 'tx2']):
                                        is_purple = True
                                elif srgb_clr is not None:
                                    val = srgb_clr.get('val')
                                    if val:
                                        clean_val = str(val).lstrip('#').upper()
                                        if len(clean_val) >= 6:
                                            clean_val = clean_val[:6]
                                            if self._is_purple_color(clean_val):
                                                is_purple = True
                                
                                if is_purple:
                                    for child in list(solid_fill):
                                        solid_fill.remove(child)
                                    new_srgb = OxmlElement('a:srgbClr')
                                    new_srgb.set('val', replacement_hex)
                                    solid_fill.append(new_srgb)
                except:
                    pass
        
        # Replace in document body (shapes, textboxes, etc.) - MULTIPLE PASSES
        # CRITICAL: Process body FIRST, then document element, then sections
        try:
            for _ in range(3):  # 3 passes to catch all shapes (increased from 2)
                replace_in_element_tree(doc.element.body)
                replace_theme_colors_aggressive(doc.element.body)
        except:
            pass
        
        # Replace in document element itself (includes all embedded content)
        try:
            for _ in range(3):  # 3 passes
                replace_in_element_tree(doc.element)
                replace_theme_colors_aggressive(doc.element)
        except:
            pass
        
        # Replace in all sections (headers, footers) - purple shapes here too
        for section in doc.sections:
            for elem in [section.header, section.footer, 
                        getattr(section, 'first_page_header', None),
                        getattr(section, 'first_page_footer', None)]:
                if elem is not None:
                    try:
                        for _ in range(3):  # 3 passes
                            replace_in_element_tree(elem._element)
                            replace_theme_colors_aggressive(elem._element)
                    except:
                        pass
        
        # FINAL PASS: Direct XML string replacement for known purple colors
        # DISABLED: This aggressive XML manipulation causes document corruption
        # The other methods (_replace_purple_in_styles, _replace_purple_in_numbering, etc.) 
        # handle color replacement safely without breaking document structure
        # self._replace_purple_in_raw_xml(doc)
    
    def _replace_purple_in_raw_xml(self, doc: Document):
        """
        Direct XML manipulation to replace ALL purple color references.
        Also replaces theme color references with MELON.
        """
        import re
        try:
            from lxml import etree
        except ImportError:
            return  # lxml not available, skip this step
        
        melon = 'FF7E51'
        
        # Known purple hex codes to replace
        purple_codes = [
            '7030A0', '800080', '660066', '9933FF', '663399',
            '8B008B', '9400D3', 'BA55D3', '9932CC', '8A2BE2', 
            'A020F0', '6A0DAD', '551A8B', '9370DB', '7B68EE', 
            'EE82EE', 'DA70D6', 'FF00FF', 'C71585', '6B3FA0', 
            '5B2C6F', '6C3483', '7D3C98', '8E44AD', '9B59B6',
        ]
        
        def replace_in_xml_element(element):
            """Replace purple colors and theme colors in element's XML - CONSERVATIVE approach."""
            try:
                # CONSERVATIVE: Only modify color attributes directly, don't rebuild elements
                # This preserves document relationships and prevents corruption
                
                # Find and modify color elements directly without rebuilding
                w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                a_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
                
                # Modify color elements directly (safer than string replacement + rebuild)
                for color_elem in element.iter(w_ns + 'color'):
                    try:
                        val = color_elem.get('val', '')
                        if val:
                            val_upper = val.upper().replace('#', '')
                            # Check if it's a purple color
                            if val_upper in purple_codes or any(p.upper() in val_upper for p in purple_codes):
                                color_elem.set('val', melon)
                    except:
                        pass
                
                # Modify drawingML color elements (for shapes)
                for srgb_elem in element.iter(a_ns + 'srgbClr'):
                    try:
                        val = srgb_elem.get('val', '')
                        if val:
                            val_upper = val.upper().replace('#', '')
                            if val_upper in purple_codes or any(p.upper() in val_upper for p in purple_codes):
                                srgb_elem.set('val', melon)
                    except:
                        pass
                
                # Replace schemeClr (theme colors) with srgbClr
                for scheme_elem in list(element.iter(a_ns + 'schemeClr')):
                    try:
                        val = scheme_elem.get('val', '')
                        if val and ('accent' in val.lower() or val.lower() in ['dk1', 'dk2', 'tx1', 'tx2', 'lt1', 'lt2']):
                            # Replace schemeClr with srgbClr
                            parent = scheme_elem.getparent()
                            if parent is not None:
                                # Create new srgbClr element
                                srgb_elem = etree.Element(a_ns + 'srgbClr')
                                srgb_elem.set('val', melon)
                                # Replace schemeClr with srgbClr
                                parent.replace(scheme_elem, srgb_elem)
                    except:
                        pass
                        
            except Exception as e:
                # Skip problematic elements to prevent document corruption
                pass
        
        # DISABLED: Raw XML manipulation on document body - too risky, causes corruption
        # The other methods (_replace_purple_text_colors, etc.) handle color replacement safely
        # Only apply to specific small elements if needed, not entire document body
        
        # Skip applying to document body - let other safer methods handle it
        # This prevents corruption of complex documents
        
        # Only apply to headers/footers if absolutely necessary (and even then, be very careful)
        # for section in doc.sections:
        #     for elem in [section.header, section.footer,
        #                 getattr(section, 'first_page_header', None),
        #                 getattr(section, 'first_page_footer', None)]:
        #         if elem is not None:
        #             try:
        #                 replace_in_xml_element(elem._element)
        #             except:
        #                 pass
    
    def _replace_purple_text_colors(self, doc: Document, replacement_color: str = 'orange'):
        """
        Find and REPLACE purple text colors with specified color.
        Processes ALL text including:
        - Regular paragraphs
        - Headings and subheadings
        - Numbers (standalone or in lists) - ALL font types
        - Numbering/list definitions
        - Document styles
        - Tables
        - Headers and footers
        - Direct XML color elements (comprehensive fallback)
        
        Args:
            doc: Word document
            replacement_color: 'orange' (default) or 'black' - color to replace purple with
        """
        # First, replace purple in numbering definitions (list numbers - all fonts)
        self._replace_purple_in_numbering(doc, replacement_color)
        
        # Replace purple in document styles (all heading/paragraph styles)
        self._replace_purple_in_styles(doc, replacement_color)
        
        # Direct XML replacement for comprehensive coverage
        # DISABLED: Can cause corruption in complex documents - use safer methods instead
        # self._replace_purple_in_document_xml(doc)
        
        # Replace purple background/shading colors (table cells, paragraphs, etc.)
        self._replace_purple_backgrounds(doc, replacement_color)
        
        # Process all paragraphs (includes headings, subheadings, numbered lists)
        # MULTIPLE PASSES to catch all purple text
        for _ in range(3):  # 3 passes for text colors (increased from 2)
            for para in doc.paragraphs:
                # CRITICAL: Check paragraph-level style for purple color
                # If paragraph style has purple, apply orange to ALL runs in the paragraph
                try:
                    para_style = para.style
                    if para_style and para_style.font and para_style.font.color:
                        if para_style.font.color.rgb:
                            color_hex = str(para_style.font.color.rgb)
                            clean_hex = color_hex.lstrip('#').upper()
                            if len(clean_hex) >= 6:
                                clean_hex = clean_hex[:6]
                                if self._is_purple_color(clean_hex):
                                    # Paragraph style is purple - apply replacement color to ALL runs
                                    replacement_rgb = self._get_replacement_color(replacement_color)
                                    for run in para.runs:
                                        run.font.color.rgb = replacement_rgb
                        elif para_style.font.color.theme_color:
                            theme = str(para_style.font.color.theme_color)
                            if 'ACCENT' in theme or 'PURPLE' in theme.upper():
                                # Paragraph style has purple theme - apply replacement color to ALL runs
                                replacement_rgb = self._get_replacement_color(replacement_color)
                                for run in para.runs:
                                    run.font.color.rgb = replacement_rgb
                except:
                    pass
                
                # Also check individual runs (for cases where runs override paragraph style)
                for run in para.runs:
                    self._replace_run_purple_color(run, replacement_color)
            
            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                self._replace_run_purple_color(run, replacement_color)
        
        # ALSO process text in shapes and textboxes (purple text in shapes)
        # This catches purple text inside purple shapes
        try:
            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            a_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
            
            # Find all textboxes and shapes with text
            for txbx in doc.element.body.iter(a_ns + 'txbx'):
                try:
                    # Get textbox content
                    for txbx_content in txbx.iter(w_ns + 'txbxContent'):
                        for para_elem in txbx_content.iter(w_ns + 'p'):
                            for r_elem in para_elem.iter(w_ns + 'r'):
                                # Check for color in run properties
                                rPr = r_elem.find(w_ns + 'rPr')
                                if rPr is not None:
                                    color_elem = rPr.find(w_ns + 'color')
                                    if color_elem is not None:
                                        val = color_elem.get('val') or color_elem.get(qn('w:val'))
                                        if val:
                                            clean_val = str(val).lstrip('#').upper()
                                            if len(clean_val) >= 6:
                                                clean_val = clean_val[:6]
                                                if self._is_purple_color(clean_val):
                                                    replacement_hex = self._get_replacement_color_hex(replacement_color)
                                                    color_elem.set('val', replacement_hex)
                                                    if color_elem.get(qn('w:val')):
                                                        color_elem.set(qn('w:val'), 'FF7E51')
                except:
                    pass
        except:
            pass
        
        # Process headers in all sections
        for section in doc.sections:
            # Main header
            try:
                for para in section.header.paragraphs:
                    for run in para.runs:
                        self._replace_run_purple_color(run)
            except:
                pass
            
            # First page header
            try:
                for para in section.first_page_header.paragraphs:
                    for run in para.runs:
                        self._replace_run_purple_color(run)
            except:
                pass
            
            # Main footer
            try:
                for para in section.footer.paragraphs:
                    for run in para.runs:
                        self._replace_run_purple_color(run)
            except:
                pass
            
            # First page footer
            try:
                for para in section.first_page_footer.paragraphs:
                    for run in para.runs:
                        self._replace_run_purple_color(run)
            except:
                pass
    
    def _replace_logos_in_document_body(self, doc: Document):
        """
        Find and REPLACE old DXC logos in document body ONLY.
        CRITICAL RULES:
        1. ONLY replaces existing logos (small images <= 10cm x 4cm)
        2. NEVER adds logos to content
        3. NEVER touches footer (footer is handled separately)
        4. Content images (large) are PRESERVED as-is
        """
        logo_path = self._find_logo()
        if not logo_path:
            return
        
        # Logo detection thresholds:
        # Old DXC logo with "DXC.technology" text can be up to 10cm wide
        # Height typically under 4cm for logos
        MAX_LOGO_WIDTH_CM = 10
        MAX_LOGO_HEIGHT_CM = 4
        
        def is_likely_logo(width_cm, height_cm, drawing=None):
            """Check if image dimensions suggest it's a logo."""
            # Logo criteria: width <= 10cm AND height <= 4cm
            # This catches the old black DXC.technology logo and purple DXC logos
            size_match = width_cm <= MAX_LOGO_WIDTH_CM and height_cm <= MAX_LOGO_HEIGHT_CM
            
            # ALSO check if it's a purple-colored image (old DXC purple logo)
            # Even if slightly larger, if it's purple, it's likely an old logo
            if drawing is not None and not size_match:
                try:
                    # Check if image has purple colors (via blip or image data)
                    # This is a heuristic - purple images in header/content are likely old logos
                    # We'll be more aggressive for purple images
                    a_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
                    blip = drawing.find('.//' + a_ns + 'blip')
                    if blip is not None:
                        # If it's a reasonable size (not huge content image), treat as logo
                        if width_cm <= 15 and height_cm <= 6:  # Slightly larger threshold for purple logos
                            return True
                except:
                    pass
            
            return size_match
        
        # Track paragraphs that had logos replaced (to add spacing after)
        logo_paragraphs = []
        
        def replace_drawing(drawing, run, para=None):
            """
            Replace a drawing element with new logo ONLY if it's a logo.
            CRITICAL: Only replaces existing logos, never adds logos to content or footer.
            """
            # SAFETY CHECK: Never process footer paragraphs
            if para is not None:
                try:
                    para_parent = para._element.getparent()
                    if para_parent is not None:
                        parent_tag = str(para_parent.tag).lower()
                        # If paragraph is in footer, skip it
                        if 'footer' in parent_tag:
                            return False
                except:
                    pass
            
            try:
                extent = drawing.find('.//' + qn('wp:extent'))
                if extent is not None:
                    cx = int(extent.get('cx', 0))  # width in EMUs
                    cy = int(extent.get('cy', 0))  # height in EMUs
                    
                    # Convert EMUs to CM (914400 EMUs = 1 inch, 1 inch = 2.54 cm)
                    width_cm = cx / 914400 * 2.54
                    height_cm = cy / 914400 * 2.54
                    
                    # ONLY replace if it's a logo (small size OR purple logo)
                    if is_likely_logo(width_cm, height_cm, drawing):
                        # Remove old logo FIRST
                        parent = drawing.getparent()
                        if parent is not None:
                            parent.remove(drawing)
                        
                        # Add new logo in SAME position
                        try:
                            new_width = Cm(min(width_cm, 4.5))
                            run.add_picture(logo_path, width=new_width)
                            
                            # Track this paragraph for spacing
                            if para is not None and para not in logo_paragraphs:
                                logo_paragraphs.append(para)
                            
                            return True
                        except:
                            pass
            except:
                pass
            return False
        
        # PRIORITY: First paragraph (header ke neeche) - REMOVE old black logo ONLY (don't add new)
        # CRITICAL: This must happen BEFORE processing other paragraphs
        if len(doc.paragraphs) > 0:
            first_para = doc.paragraphs[0]
            # Verify this is actually a body paragraph, not accidentally a footer
            try:
                para_parent = first_para._element.getparent()
                if para_parent is not None:
                    parent_tag = str(para_parent.tag).lower()
                    if 'footer' in parent_tag or 'header' in parent_tag:
                        # Skip if it's in header/footer
                        pass
                    else:
                        # Process first paragraph for logo REMOVAL ONLY (no replacement)
                        for run in list(first_para.runs):
                            run_element = run._element
                            # Find all drawings (old logos) and REMOVE them (don't replace)
                            drawings = list(run_element.findall('.//' + qn('w:drawing')))
                            for drawing in drawings:
                                try:
                                    # Check if it's a logo-sized image
                                    extent = drawing.find('.//' + qn('wp:extent'))
                                    if extent is not None:
                                        cx = int(extent.get('cx', 0))
                                        cy = int(extent.get('cy', 0))
                                        if cx > 0 and cy > 0:
                                            width_cm = cx / 914400 * 2.54
                                            height_cm = cy / 914400 * 2.54
                                            # If it's logo-sized, REMOVE it (don't replace)
                                            if width_cm <= 15 and height_cm <= 6:
                                                parent = drawing.getparent()
                                                if parent is not None:
                                                    parent.remove(drawing)
                                except:
                                    pass
                            
                            # Check for blip (embedded images) - REMOVE them
                            blips = list(run_element.findall('.//' + qn('a:blip')))
                            for blip in blips:
                                try:
                                    parent = blip.getparent()
                                    while parent is not None and not parent.tag.endswith('drawing'):
                                        parent = parent.getparent()
                                    if parent is not None:
                                        # Remove the drawing (don't replace)
                                        drawing_parent = parent.getparent()
                                        if drawing_parent is not None:
                                            drawing_parent.remove(parent)
                                except:
                                    pass
            except:
                pass
        
        # Search through all paragraphs in document body ONLY
        # CRITICAL: Only process body paragraphs, NEVER footer
        # doc.paragraphs contains ONLY body paragraphs (not header/footer)
        # BUT: We need to verify each paragraph is actually in body, not accidentally in footer
        
        # Get body element to verify paragraphs belong to it
        body_element = doc.element.body
        
        for para in doc.paragraphs:
            # TRIPLE CHECK: Ensure this paragraph is in body, not footer
            try:
                para_elem = para._element
                para_parent = para_elem.getparent()
                
                # Check 1: Parent tag should be body
                if para_parent is not None:
                    parent_tag = str(para_parent.tag).lower()
                    # Skip if paragraph is in footer or header
                    if 'footer' in parent_tag or 'header' in parent_tag:
                        continue
                
                # Check 2: Verify paragraph is actually a child of body element
                # Walk up the tree to find body
                current = para_elem
                found_body = False
                while current is not None:
                    if current == body_element:
                        found_body = True
                        break
                    current = current.getparent()
                
                if not found_body:
                    continue  # Not in body, skip it
                    
            except:
                # If we can't verify, skip to be safe
                continue
            
            for run in list(para.runs):
                run_element = run._element
                # Find all drawings (images) - only replace if they're logos
                drawings = list(run_element.findall('.//' + qn('w:drawing')))
                for drawing in drawings:
                    # Only replace if it's a logo (small size), not content images
                    replace_drawing(drawing, run, para)
                
                # Check for blip (embedded images)
                blips = list(run_element.findall('.//' + qn('a:blip')))
                for blip in blips:
                    try:
                        parent = blip.getparent()
                        while parent is not None and not parent.tag.endswith('drawing'):
                            parent = parent.getparent()
                        if parent is not None:
                            # Only replace if it's a logo
                            replace_drawing(parent, run, para)
                    except:
                        pass
        
        # Also check tables for logos (only REPLACE existing logos in tables)
        # Tables are always in body, so safe to process
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        # Verify paragraph is in table cell (body), not footer
                        try:
                            para_parent = para._element.getparent()
                            if para_parent is not None:
                                parent_tag = str(para_parent.tag).lower()
                                if 'footer' in parent_tag:
                                    continue
                        except:
                            pass
                        
                        for run in list(para.runs):
                            run_element = run._element
                            drawings = list(run_element.findall('.//' + qn('w:drawing')))
                            for drawing in drawings:
                                # Only replace if it's a logo (not content images)
                                replace_drawing(drawing, run, para)
        
        # NO XML scan - we've already processed all body paragraphs and tables
        # XML scan might accidentally pick up footer elements, so we skip it completely
        
        # Add proper spacing after logo paragraphs using XML directly
        # This ensures spacing is actually applied
        for para in logo_paragraphs:
            try:
                w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                p_elem = para._element
                
                # Find or create paragraph properties
                pPr = p_elem.find(w_ns + 'pPr')
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    p_elem.insert(0, pPr)
                
                # Find or create spacing element
                spacing = pPr.find(w_ns + 'spacing')
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                
                # Set spacing after: 480 twips = 24pt (1 twip = 1/20 pt)
                # Using 600 twips = 30pt for more visible spacing
                spacing.set(qn('w:after'), '600')
                spacing.set(qn('w:before'), '120')  # 6pt before
                
            except Exception:
                # Fallback to python-docx method
                try:
                    para.paragraph_format.space_after = Pt(30)
                    para.paragraph_format.space_before = Pt(6)
                except:
                    pass
    
    def _replace_header_content(self, header, logo_path):
        """
        Replace header: Remove ALL old logos FIRST, then add ONE new logo.
        """
        if header is None:
            return
            
        header.is_linked_to_previous = False
        header_element = header._element
        
        # STEP 1: Remove ALL old logos/images from header paragraphs
        # CRITICAL: Remove ALL images, regardless of size or color (header should only have new logo)
        try:
            for para in list(header.paragraphs):
                for run in list(para.runs):
                    run_element = run._element
                    # Find and remove ALL drawings (images/logos) - AGGRESSIVE
                    drawings = list(run_element.findall('.//' + qn('w:drawing')))
                    for drawing in drawings:
                        try:
                            parent = drawing.getparent()
                            if parent is not None:
                                parent.remove(drawing)
                        except:
                            pass
                    # Also remove blip elements (image references)
                    a_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
                    blips = list(run_element.findall('.//' + a_ns + 'blip'))
                    for blip in blips:
                        try:
                            # Find the drawing parent and remove it
                            parent = blip.getparent()
                            while parent is not None and not parent.tag.endswith('drawing'):
                                parent = parent.getparent()
                            if parent is not None:
                                parent_elem = parent.getparent()
                                if parent_elem is not None:
                                    parent_elem.remove(parent)
                        except:
                            pass
        except:
            pass
        
        # STEP 2: Remove ALL drawings from header XML directly (AGGRESSIVE)
        try:
            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            wp_ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
            a_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
            
            # Remove ALL drawings (multiple iterations to catch nested ones)
            for _ in range(3):  # Multiple passes to catch all
                for drawing in list(header_element.iter(wp_ns + 'drawing')):
                    try:
                        parent = drawing.getparent()
                        if parent is not None:
                            parent.remove(drawing)
                    except:
                        pass
                for drawing in list(header_element.iter(w_ns + 'drawing')):
                    try:
                        parent = drawing.getparent()
                        if parent is not None:
                            parent.remove(drawing)
                    except:
                        pass
                # Also remove blip elements (image references)
                for blip in list(header_element.iter(a_ns + 'blip')):
                    try:
                        parent = blip.getparent()
                        while parent is not None and not parent.tag.endswith('drawing'):
                            parent = parent.getparent()
                        if parent is not None:
                            parent_elem = parent.getparent()
                            if parent_elem is not None:
                                parent_elem.remove(parent)
                    except:
                        pass
        except:
            pass
        
        # STEP 3: Clear ALL header content completely
        for child in list(header_element):
            try:
                header_element.remove(child)
            except:
                pass
        
        # STEP 4: Final aggressive check - remove ANY remaining drawings/images
        try:
            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            wp_ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
            a_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
            
            # Final pass - remove any remaining drawings
            for drawing in list(header_element.iter(wp_ns + 'drawing')):
                try:
                    parent = drawing.getparent()
                    if parent is not None:
                        parent.remove(drawing)
                except:
                    pass
            for drawing in list(header_element.iter(w_ns + 'drawing')):
                try:
                    parent = drawing.getparent()
                    if parent is not None:
                        parent.remove(drawing)
                except:
                    pass
            # Remove any remaining blip elements
            for blip in list(header_element.iter(a_ns + 'blip')):
                try:
                    parent = blip.getparent()
                    while parent is not None and not parent.tag.endswith('drawing'):
                        parent = parent.getparent()
                    if parent is not None:
                        parent_elem = parent.getparent()
                        if parent_elem is not None:
                            parent_elem.remove(parent)
                except:
                    pass
        except:
            pass
        
        # STEP 5: Add ONE new logo only (after all old logos are removed)
        if logo_path:
            try:
                header_para = header.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = header_para.add_run()
                run.add_picture(logo_path, width=Cm(4.5))
            except:
                pass
    
    def _apply_dxc_header(self, doc: Document):
        """
        REPLACE old header with new DXC branded header.
        Handles both regular headers AND first page headers.
        ONLY touches the header area - document body is NEVER modified.
        """
        logo_path = self._find_logo()
        
        for section in doc.sections:
            # Replace the main/default header
            header = section.header
            self._replace_header_content(header, logo_path)
            
            # Also replace FIRST PAGE header if document has different first page
            # This ensures old logo on first page header also gets replaced
            if section.different_first_page_header_footer:
                first_page_header = section.first_page_header
                self._replace_header_content(first_page_header, logo_path)
            else:
                # Check if first page header has content (some docs have it even without the flag)
                try:
                    first_page_header = section.first_page_header
                    if first_page_header._element is not None and len(list(first_page_header._element)) > 0:
                        self._replace_header_content(first_page_header, logo_path)
                except:
                    pass
    
    def _replace_footer_content(self, footer):
        """
        Footer: Remove ALL old logos, clear completely. Footer stays EMPTY - NO LOGOS ADDED.
        CRITICAL: Must remove drawings from WITHIN runs, not just paragraphs.
        """
        if footer is None:
            return
            
        footer.is_linked_to_previous = False
        footer_element = footer._element
        
        # STEP 1: Remove ALL drawings from WITHIN runs FIRST (before removing paragraphs)
        # This is critical - drawings are nested inside runs
        # MULTIPLE PASSES to ensure all drawings are removed
        for pass_num in range(3):  # 3 passes to catch all drawings
            try:
                w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                wp_ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
                a_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
                
                # Find ALL runs in footer
                for run_elem in list(footer_element.iter(w_ns + 'r')):
                    # Find ALL drawings within this run
                    for drawing in list(run_elem.iter(wp_ns + 'drawing')):
                        try:
                            parent = drawing.getparent()
                            if parent is not None:
                                parent.remove(drawing)
                        except:
                            pass
                    
                    # Also remove blips (image references) from runs
                    for blip in list(run_elem.iter(a_ns + 'blip')):
                        try:
                            # Find the drawing parent
                            parent = blip.getparent()
                            while parent is not None and not parent.tag.endswith('drawing'):
                                parent = parent.getparent()
                            if parent is not None:
                                drawing_parent = parent.getparent()
                                if drawing_parent is not None:
                                    drawing_parent.remove(parent)
                        except:
                            pass
            except:
                pass
        
        # STEP 2: Remove ALL paragraphs (AGGRESSIVE - remove all child elements first)
        try:
            # Multiple passes to ensure all paragraphs are removed
            for _ in range(3):
                for para in list(footer.paragraphs):
                    try:
                        para_elem = para._element
                        # Remove ALL child elements from paragraph (runs, drawings, etc.)
                        for child in list(para_elem):
                            try:
                                para_elem.remove(child)
                            except:
                                pass
                        # Then remove paragraph itself
                        footer_element.remove(para_elem)
                    except:
                        pass
        except:
            pass
        
        # STEP 3: Remove ALL drawings/images from footer XML (aggressive - catch any remaining)
        try:
            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            wp_ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
            a_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
            
            # Multiple passes to catch all drawings
            for _ in range(3):
                # Remove all drawings (any namespace)
                for drawing in list(footer_element.iter(wp_ns + 'drawing')):
                    try:
                        parent = drawing.getparent()
                        if parent is not None:
                            parent.remove(drawing)
                    except:
                        pass
                
                # Remove all blip elements (image references)
                for blip in list(footer_element.iter(a_ns + 'blip')):
                    try:
                        parent = blip.getparent()
                        while parent is not None and not parent.tag.endswith('drawing'):
                            parent = parent.getparent()
                        if parent is not None:
                            parent_elem = parent.getparent()
                            if parent_elem is not None:
                                parent_elem.remove(parent)
                    except:
                        pass
        except:
            pass
        
        # STEP 4: Clear ALL child elements - footer completely empty
        for child in list(footer_element):
            try:
                footer_element.remove(child)
            except:
                pass
        
        # STEP 5: Final aggressive cleanup - remove ANY remaining content
        try:
            # Remove any remaining paragraphs
            while len(footer.paragraphs) > 0:
                try:
                    para = footer.paragraphs[0]
                    # Remove all runs first
                    for run in list(para.runs):
                        try:
                            para._element.remove(run._element)
                        except:
                            pass
                    footer_element.remove(para._element)
                except:
                    break
            
            # Remove any remaining elements
            for elem in list(footer_element):
                try:
                    footer_element.remove(elem)
                except:
                    pass
        except:
            pass
        
        # Footer is now COMPLETELY EMPTY - NO paragraphs, NO content, NO logos
        # DO NOT ADD ANYTHING TO FOOTER
    
    def _apply_dxc_footer(self, doc: Document):
        """
        REPLACE old footer with new DXC branded footer.
        Handles both regular footers AND first page footers.
        ONLY touches the footer area - document body is NEVER modified.
        CRITICAL: Footer must be COMPLETELY EMPTY - NO logos, NO content.
        """
        # Process ALL sections - clear ALL footers (MULTIPLE PASSES)
        # Pass 1: Initial clearing
        for section in doc.sections:
            # Clear main/default footer
            try:
                footer = section.footer
                if footer is not None:
                    self._replace_footer_content(footer)
            except:
                pass
            
            # Clear first page footer (always, regardless of flag)
            try:
                first_page_footer = section.first_page_footer
                if first_page_footer is not None:
                    self._replace_footer_content(first_page_footer)
            except:
                pass
        
        # Pass 2: Aggressive clearing (in case something was added)
        for section in doc.sections:
            try:
                footer = section.footer
                if footer is not None:
                    self._replace_footer_content(footer)
            except:
                pass
            
            try:
                first_page_footer = section.first_page_footer
                if first_page_footer is not None:
                    self._replace_footer_content(first_page_footer)
            except:
                pass
        
        # Pass 3: Final verification pass - ensure footers are completely empty
        for section in doc.sections:
            try:
                footer = section.footer
                if footer is not None:
                    # Check if footer still has content
                    if len(footer.paragraphs) > 0:
                        self._replace_footer_content(footer)
            except:
                pass
            
            try:
                first_page_footer = section.first_page_footer
                if first_page_footer is not None:
                    if len(first_page_footer.paragraphs) > 0:
                        self._replace_footer_content(first_page_footer)
            except:
                pass
    
    def rebrand_document(self, source_file: bytes, source_filename: str, replacement_color: str = 'orange') -> Tuple[bytes, str]:
        """
        Rebrand a Word document with DXC branding.
        
        - REPLACES old DXC logos anywhere in document with new logo
        - REPLACES purple text colors with specified color (orange or black)
        
        Args:
            source_file: Word document file bytes
            source_filename: Original filename
            replacement_color: 'orange' (default) or 'black' - color to replace purple shapes with
        - REPLACES header with new DXC header
        - REPLACES footer with new DXC footer
        - All other content (text, tables, fonts, formatting) PRESERVED EXACTLY
        """
        
        # Save source to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(source_file)
            source_path = tmp.name
        
        try:
            # Load the ORIGINAL document - we will modify it in place
            doc = Document(source_path)
            
            # Update document properties
            doc.core_properties.author = "DXC Technology"
            doc.core_properties.company = "DXC Technology"
            doc.core_properties.modified = datetime.now()
            
            # CRITICAL ORDER: Process in this exact sequence
            
            # STEP 1: REPLACE purple colors FIRST
            # This converts purple shapes, text, tables to specified color
            self._replace_purple_text_colors(doc, replacement_color)
            
            # STEP 2: REMOVE old logos from header, then add new logo
            # Header gets ONE new logo only
            self._apply_dxc_header(doc)
            
            # STEP 3: REPLACE old DXC logos in document body (first paragraph + rest)
            # This handles old black logo below header
            # CRITICAL: This only processes body paragraphs, never touches footer
            self._replace_logos_in_document_body(doc)
            
            # STEP 4: CLEAR footer LAST (after everything else)
            # This ensures footer is empty even if something added logos to it
            # Footer must be completely empty - NO logos
            self._apply_dxc_footer(doc)
            
            # STEP 5: FINAL PASS - Clear footer ONE MORE TIME to ensure it's empty
            # This is critical - sometimes logos get added during processing
            self._apply_dxc_footer(doc)
            
            # Save the modified document
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            # Generate output filename
            output_name = f"DXC_Rebranded_{Path(source_filename).stem}_{datetime.now().strftime('%Y%m%d')}.docx"
            
            return doc_bytes.getvalue(), output_name
            
        finally:
            # Cleanup
            if os.path.exists(source_path):
                os.unlink(source_path)


# ============================================================================
# UI/UX CODE (Same as existing files)
# ============================================================================

# Page Configuration
try:
    st.set_page_config(
        page_title="Reface | DXC Technology",
        page_icon=None,
        layout="wide",
        initial_sidebar_state="expanded"
    )
except st.errors.StreamlitAPIException:
    pass


def get_logo_base64():
    """Load DXC logo as base64"""
    logo_path = Path(__file__).parent / "dxc_brand_logo.png"
    if logo_path.exists():
        with open(logo_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return None


def load_css():
    """Load enterprise CSS styling - EXACT SAME as existing files"""
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Source+Sans+Pro:wght@300;400;500;600;700&display=swap');
    
    :root {
        --dxc-dark: #0E1020;
        --dxc-orange: #EF8900;
        --dxc-gradient-start: #5B8DEF;
        --dxc-gradient-end: #EF8068;
        --dxc-light-bg: #F8F9FA;
        --success-green: #198754;
        --error-red: #DC3545;
        --warning-amber: #FFC107;
        --text-primary: #212529;
        --text-secondary: #6C757D;
        --bg-white: #FFFFFF;
        --border-light: #DEE2E6;
    }
    
    /* Hide Streamlit branding and unnecessary elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stDeployButton {display: none;}
    
    /* Remove default Streamlit container padding */
    .block-container {
        padding-top: 5rem;
        padding-bottom: 5rem;
    }
    
    /* Main app styling */
    .main .block-container {
        max-width: 1200px;
        margin: 0 auto;
    }
    
    /* Enterprise Header - Fixed to top */
    .enterprise-header {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        z-index: 999;
        background: linear-gradient(135deg, var(--dxc-dark) 0%, #1C1E30 50%, #2A2D45 100%);
        padding: 0.75rem 1rem;
        border-radius: 0;
        margin: 0;
        display: flex;
        align-items: center;
        justify-content: space-between;
        border-bottom: 2px solid var(--dxc-orange);
    }
    
    .header-left {
        display: flex;
        align-items: center;
        gap: 0.75rem;
    }
    
    .header-logo img {
        height: 36px;
        width: auto;
    }
    
    .header-content {
        border-left: 1px solid rgba(255,255,255,0.2);
        padding-left: 0.75rem;
    }
    
    .enterprise-title {
        font-family: 'Source Sans Pro', sans-serif;
        font-size: 1.25rem;
        font-weight: 600;
        color: #FFFFFF;
        margin: 0;
        letter-spacing: -0.5px;
    }
    
    .enterprise-subtitle {
        font-family: 'Source Sans Pro', sans-serif;
        font-size: 0.75rem;
        color: rgba(255,255,255,0.7);
        margin-top: 0.1rem;
        font-weight: 400;
    }
    
    /* Section Headers */
    .section-header {
        font-family: 'Source Sans Pro', sans-serif;
        font-size: 1.1rem;
        font-weight: 600;
        color: var(--dxc-dark);
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid var(--dxc-orange);
        display: flex;
        align-items: center;
        gap: 0.5rem;
    }
    
    /* Status indicators */
    .status-indicator {
        display: inline-flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.5rem 1rem;
        border-radius: 4px;
        font-size: 0.875rem;
        font-weight: 500;
        margin: 0.25rem 0;
    }
    
    .status-ready {
        background: rgba(25, 135, 84, 0.1);
        color: var(--success-green);
        border: 1px solid rgba(25, 135, 84, 0.3);
    }
    
    .status-unavailable {
        background: rgba(220, 53, 69, 0.1);
        color: var(--error-red);
        border: 1px solid rgba(220, 53, 69, 0.3);
    }
    
    .status-dot {
        width: 8px;
        height: 8px;
        border-radius: 50%;
        display: inline-block;
    }
    
    .status-dot.green {
        background: var(--success-green);
    }
    
    .status-dot.red {
        background: var(--error-red);
    }
    
    /* File upload success message */
    .upload-success {
        background: linear-gradient(135deg, rgba(25, 135, 84, 0.08) 0%, rgba(25, 135, 84, 0.04) 100%);
        border: 1px solid rgba(25, 135, 84, 0.2);
        border-left: 4px solid var(--success-green);
        color: #0F5132;
        padding: 1rem 1.25rem;
        border-radius: 4px;
        font-weight: 500;
        margin: 1rem 0;
    }
    
    /* Error message */
    .upload-error {
        background: linear-gradient(135deg, rgba(220, 53, 69, 0.08) 0%, rgba(220, 53, 69, 0.04) 100%);
        border: 1px solid rgba(220, 53, 69, 0.2);
        border-left: 4px solid var(--error-red);
        color: #842029;
        padding: 1rem 1.25rem;
        border-radius: 4px;
        font-weight: 500;
        margin: 0.5rem 0;
    }
    
    /* Feature cards grid */
    .feature-grid {
        display: grid;
        grid-template-columns: repeat(4, 1fr);
        gap: 1rem;
        margin: 1.5rem 0;
    }
    
    .feature-card {
        background: var(--bg-white);
        padding: 1.25rem;
        border-radius: 8px;
        border: 1px solid var(--border-light);
        text-align: center;
        transition: all 0.2s ease;
    }
    
    .feature-card:hover {
        border-color: var(--dxc-orange);
        box-shadow: 0 4px 12px rgba(14, 16, 32, 0.08);
    }
    
    .feature-icon {
        width: 40px;
        height: 40px;
        margin: 0 auto 0.75rem;
        background: linear-gradient(135deg, var(--dxc-gradient-start), var(--dxc-gradient-end));
        border-radius: 8px;
        display: flex;
        align-items: center;
        justify-content: center;
    }
    
    .feature-icon svg {
        width: 20px;
        height: 20px;
        fill: white;
    }
    
    .feature-title {
        font-weight: 600;
        color: var(--dxc-dark);
        font-size: 0.9rem;
        margin-bottom: 0.25rem;
    }
    
    .feature-desc {
        font-size: 0.8rem;
        color: var(--text-secondary);
    }
    
    /* Stats cards */
    .stats-row {
        display: flex;
        gap: 1rem;
        margin: 1rem 0;
    }
    
    .stat-card {
        flex: 1;
        background: var(--bg-white);
        border: 1px solid var(--border-light);
        border-radius: 8px;
        padding: 1.25rem;
        text-align: center;
    }
    
    .stat-value {
        font-size: 2rem;
        font-weight: 700;
        color: var(--dxc-dark);
        line-height: 1;
    }
    
    .stat-label {
        font-size: 0.85rem;
        color: var(--text-secondary);
        margin-top: 0.5rem;
    }
    
    /* Sidebar styling */
    section[data-testid="stSidebar"] {
        background: #F6F3F0;
        border-right: none;
        box-shadow: 2px 0 8px rgba(14, 16, 32, 0.06);
        z-index: auto !important;
    }
    
    section[data-testid="stSidebar"] > div:first-child {
        padding-top: 1.5rem;
    }
    
    section[data-testid="stSidebar"] .section-header {
        font-size: 0.85rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        color: var(--dxc-dark);
        margin-bottom: 0.75rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid var(--dxc-orange);
    }
    
    /* Sidebar section cards */
    .sidebar-section {
        background: #FFFFFF;
        border-radius: 8px;
        padding: 1rem;
        margin-bottom: 1rem;
        border: 1px solid rgba(14, 16, 32, 0.08);
        box-shadow: 0 1px 3px rgba(14, 16, 32, 0.04);
    }
    
    .sidebar-section-title {
        font-size: 0.8rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.5px;
        color: var(--dxc-dark);
        margin-bottom: 0.75rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid var(--dxc-orange);
    }
    
    /* Sidebar status pills */
    .sidebar-status {
        display: flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.6rem 0.75rem;
        border-radius: 6px;
        font-size: 0.8rem;
        font-weight: 500;
        margin-bottom: 0.5rem;
    }
    
    .sidebar-status.ready {
        background: linear-gradient(135deg, rgba(25, 135, 84, 0.12) 0%, rgba(25, 135, 84, 0.06) 100%);
        color: #0D6E3F;
        border: 1px solid rgba(25, 135, 84, 0.2);
    }
    
    .sidebar-status.unavailable {
        background: linear-gradient(135deg, rgba(220, 53, 69, 0.12) 0%, rgba(220, 53, 69, 0.06) 100%);
        color: #A52834;
        border: 1px solid rgba(220, 53, 69, 0.2);
    }
    
    .sidebar-status-dot {
        width: 6px;
        height: 6px;
        border-radius: 50%;
        flex-shrink: 0;
    }
    
    .sidebar-status-dot.green {
        background: var(--success-green);
        box-shadow: 0 0 4px rgba(25, 135, 84, 0.5);
    }
    
    .sidebar-status-dot.red {
        background: var(--error-red);
        box-shadow: 0 0 4px rgba(220, 53, 69, 0.5);
    }
    
    /* Sidebar list items */
    .sidebar-list {
        list-style: none;
        padding: 0;
        margin: 0;
    }
    
    .sidebar-list-item {
        display: flex;
        align-items: flex-start;
        gap: 0.5rem;
        padding: 0.4rem 0;
        font-size: 0.82rem;
        color: var(--text-primary);
        line-height: 1.4;
    }
    
    .sidebar-list-item::before {
        content: "";
        width: 4px;
        height: 4px;
        background: var(--dxc-orange);
        border-radius: 50%;
        margin-top: 0.45rem;
        flex-shrink: 0;
    }
    
    .sidebar-list-item.check::before {
        content: "✓";
        width: auto;
        height: auto;
        background: none;
        color: var(--success-green);
        font-weight: 700;
        font-size: 0.75rem;
        margin-top: 0;
    }
    
    .sidebar-list-item.update::before {
        content: "→";
        width: auto;
        height: auto;
        background: none;
        color: var(--dxc-orange);
        font-weight: 700;
        font-size: 0.75rem;
        margin-top: 0;
    }
    
    /* Sidebar template info */
    .sidebar-template {
        display: flex;
        align-items: center;
        gap: 0.5rem;
        padding: 0.5rem 0;
        font-size: 0.82rem;
    }
    
    .sidebar-template-label {
        font-weight: 600;
        color: var(--text-secondary);
        min-width: 45px;
    }
    
    .sidebar-template-value {
        color: var(--dxc-dark);
        font-weight: 500;
    }
    
    /* Sidebar divider */
    .sidebar-divider {
        height: 1px;
        background: linear-gradient(90deg, transparent, rgba(14, 16, 32, 0.1), transparent);
        margin: 1rem 0;
    }
    
    /* File list styling */
    .file-item {
        display: flex;
        align-items: center;
        padding: 0.5rem 0;
        border-bottom: 1px solid var(--border-light);
        font-size: 0.9rem;
    }
    
    .file-item:last-child {
        border-bottom: none;
    }
    
    .file-type-badge {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 32px;
        height: 32px;
        border-radius: 6px;
        font-size: 0.7rem;
        font-weight: 700;
        margin-right: 0.75rem;
        color: white;
    }
    
    .file-type-badge.docx {
        background: linear-gradient(135deg, #2B579A, #1E3A5F);
    }
    
    .file-type-badge.pdf {
        background: linear-gradient(135deg, #DC143C, #B22222);
    }
    
    /* Footer - Fixed to bottom */
    .enterprise-footer {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        z-index: 999;
        text-align: center;
    }
    
    .footer-top {
        background: #000000;
        border-top: none;
        padding: 0.4rem 1rem;
    }
    
    .footer-credits {
        display: flex;
        justify-content: center;
        align-items: center;
        gap: 0.5rem;
        flex-wrap: wrap;
        font-size: 0.7rem;
        color: rgba(255,255,255,0.8);
    }
    
    .footer-credits strong {
        color: #ffffff;
        font-weight: 600;
    }
    
    .footer-separator {
        color: rgba(255,255,255,0.5);
    }
    
    .footer-bottom {
        background: #000000;
        padding: 0.25rem 1rem;
    }
    
    .footer-copyright {
        font-size: 0.65rem;
        color: rgba(255,255,255,0.7);
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    
    .footer-developer a {
        color: var(--dxc-orange);
        text-decoration: none;
    }
    
    
    /* Button styling overrides */
    .stButton > button {
        font-family: 'Source Sans Pro', sans-serif;
        font-weight: 600;
        border-radius: 6px;
        padding: 0.5rem 1.5rem;
        transition: all 0.2s ease;
    }
    
    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, var(--dxc-dark), #1C1E30);
        border: none;
    }
    
    .stButton > button[kind="primary"]:hover {
        background: linear-gradient(135deg, #1C1E30, var(--dxc-dark));
        box-shadow: 0 4px 12px rgba(14, 16, 32, 0.3);
    }
    
    /* Download button styling */
    .stDownloadButton > button {
        font-family: 'Source Sans Pro', sans-serif;
        font-weight: 600;
        border-radius: 6px;
    }
    
    /* Progress bar styling */
    .stProgress > div > div {
        background: linear-gradient(90deg, var(--dxc-dark), var(--dxc-orange));
    }
    
    /* Metric styling */
    [data-testid="stMetric"] {
        background: var(--bg-white);
        border: 1px solid var(--border-light);
        border-radius: 8px;
        padding: 1rem;
    }
    
    [data-testid="stMetricLabel"] {
        font-size: 0.85rem;
        color: var(--text-secondary);
    }
    
    [data-testid="stMetricValue"] {
        font-size: 1.75rem;
        font-weight: 700;
        color: var(--dxc-dark);
    }
    
    /* File uploader styling */
    [data-testid="stFileUploader"] {
        border: 2px dashed var(--border-light);
        border-radius: 8px;
        padding: 1rem;
        background: var(--bg-white);
    }
    
    [data-testid="stFileUploader"]:hover {
        border-color: var(--dxc-orange);
    }
    
    /* Remove expander arrows and streamlit decorations */
    .streamlit-expanderHeader {
        font-weight: 600;
    }
    
    /* Info/warning/error boxes - cleaner look */
    .stAlert {
        border-radius: 6px;
    }
    
    /* Divider styling */
    hr {
        border: none;
        border-top: 1px solid var(--border-light);
        margin: 1.5rem 0;
    }
    </style>
    """, unsafe_allow_html=True)


# ============================================================================
# UNIFIED BULK REBRANDER
# ============================================================================

class UnifiedBulkRebrander:
    """Handles bulk rebranding of multiple PDF and Word documents"""
    
    def __init__(self):
        self.pdf_rebrander = UnifiedPDFRebrander() if PDF_ALL_DEPS_AVAILABLE else None
        self.word_rebrander = UnifiedWordRebrander() if DOCX_AVAILABLE else None
        self.results = []
        self.progress_callback = None
    
    def set_progress_callback(self, callback):
        """Set progress callback function"""
        self.progress_callback = callback
    
    def update_progress(self, message: str, progress: int, file_name: str = ""):
        """Update progress"""
        if self.progress_callback:
            self.progress_callback(message, progress, file_name)
    
    def process_files(self, uploaded_files: List, replacement_color: str = 'orange') -> Dict:
        """
        Process multiple uploaded files for rebranding (PDF and Word together).
        
        Args:
            uploaded_files: List of uploaded file objects
            replacement_color: 'orange' (default) or 'black' - color to replace purple shapes with
        
        Returns dict with:
        - successful: list of (filename, bytes, new_filename)
        - failed: list of (filename, error_message)
        - stats: processing statistics
        """
        results = {
            'successful': [],
            'failed': [],
            'stats': {
                'total': len(uploaded_files),
                'pdf_count': 0,
                'word_count': 0,
                'success_count': 0,
                'fail_count': 0,
                'processing_time': 0
            }
        }
        
        start_time = time.time()
        
        for idx, uploaded_file in enumerate(uploaded_files):
            file_name = uploaded_file.name
            file_ext = Path(file_name).suffix.lower()
            
            progress = int(((idx + 1) / len(uploaded_files)) * 100)
            self.update_progress(f"Processing {file_name}...", progress, file_name)
            
            try:
                file_bytes = uploaded_file.read()
                uploaded_file.seek(0)
                
                if file_ext == '.pdf':
                    results['stats']['pdf_count'] += 1
                    
                    if not self.pdf_rebrander:
                        results['failed'].append((file_name, "PDF rebranding not available"))
                        results['stats']['fail_count'] += 1
                        continue
                    
                    self.update_progress(f"Rebranding {file_name}...", progress, file_name)
                    
                    rebranded_bytes, new_filename = self.pdf_rebrander.rebrand_pdf(
                        file_bytes, file_name, replacement_color
                    )
                    
                    self.update_progress(f"Completed {file_name}", progress, file_name)
                    results['successful'].append((file_name, rebranded_bytes, new_filename))
                    results['stats']['success_count'] += 1
                    
                elif file_ext == '.docx':
                    results['stats']['word_count'] += 1
                    
                    if not self.word_rebrander:
                        results['failed'].append((file_name, "Word rebranding not available"))
                        results['stats']['fail_count'] += 1
                        continue
                    
                    self.update_progress(f"Rebranding {file_name}...", progress, file_name)
                    
                    rebranded_bytes, new_filename = self.word_rebrander.rebrand_document(
                        file_bytes, file_name, replacement_color
                    )
                    
                    self.update_progress(f"Completed {file_name}", progress, file_name)
                    results['successful'].append((file_name, rebranded_bytes, new_filename))
                    results['stats']['success_count'] += 1
                    
                else:
                    results['failed'].append((file_name, f"Unsupported file type: {file_ext}"))
                    results['stats']['fail_count'] += 1
                    
            except Exception as e:
                results['failed'].append((file_name, str(e)))
                results['stats']['fail_count'] += 1
        
        results['stats']['processing_time'] = round(time.time() - start_time, 2)
        
        self.update_progress("Processing complete!", 100, "")
        return results
    
    def create_zip_download(self, successful_files: List[Tuple]) -> bytes:
        """Create a ZIP file containing all successfully rebranded documents"""
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for original_name, file_bytes, new_name in successful_files:
                zip_file.writestr(new_name, file_bytes)
        
        zip_buffer.seek(0)
        return zip_buffer.getvalue()


# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    load_css()
    
    # Get logo
    logo_base64 = get_logo_base64()
    
    # Header with logo
    logo_html = f'<img src="data:image/png;base64,{logo_base64}" alt="DXC Technology">' if logo_base64 else '<span style="color: white; font-weight: 700; font-size: 1.5rem;">DXC</span>'
    
    st.markdown(f"""
    <div class="enterprise-header">
        <div class="header-left">
            <div class="header-logo">
                {logo_html}
            </div>
            <div class="header-content">
                <div class="enterprise-title">Reface</div>
                <div class="enterprise-subtitle">Brand transformation at the speed of scale.</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Check dependencies
    if not DOCX_AVAILABLE and not PDF_ALL_DEPS_AVAILABLE:
        st.markdown("""
        <div class="upload-error">
            <strong>Configuration Error:</strong> No rebranding modules available. 
            Please install required dependencies:
            <br><br>
            <strong>For Word:</strong> pip install python-docx
            <br>
            <strong>For PDF:</strong> pip install PyMuPDF PyPDF2 reportlab Pillow numpy
        </div>
        """, unsafe_allow_html=True)
        st.stop()
    
    # Feature cards
    st.markdown("""
    <div class="feature-grid" style="grid-template-columns: repeat(4, 1fr);">
        <div class="feature-card">
            <div class="feature-icon">
                <svg viewBox="0 0 24 24" fill="white"><path d="M14,2H6A2,2 0 0,0 4,4V20A2,2 0 0,0 6,22H18A2,2 0 0,0 20,20V8L14,2M18,20H6V4H13V9H18V20Z"/></svg>
            </div>
            <div class="feature-title">Word Documents</div>
            <div class="feature-desc">Rebrand .docx files</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">
                <svg viewBox="0 0 24 24" fill="white"><path d="M14,2H6A2,2 0 0,0 4,4V20A2,2 0 0,0 6,22H18A2,2 0 0,0 20,20V8L14,2M18,20H6V4H13V9H18V20Z"/></svg>
            </div>
            <div class="feature-title">PDF Documents</div>
            <div class="feature-desc">Rebrand .pdf files</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">
                <svg viewBox="0 0 24 24" fill="white"><path d="M13,2.03V2.05L13,4.05C17.39,4.59 20.5,8.58 19.96,12.97C19.5,16.61 16.64,19.5 13,19.93V21.93C18.5,21.38 22.5,16.5 21.95,11C21.5,6.25 17.73,2.5 13,2.03M11,2.06C9.05,2.25 7.19,3 5.67,4.26L7.1,5.74C8.22,4.84 9.57,4.26 11,4.06V2.06M4.26,5.67C3,7.19 2.25,9.04 2.05,11H4.05C4.24,9.58 4.8,8.23 5.69,7.1L4.26,5.67M2.06,13C2.26,14.96 3.03,16.81 4.27,18.33L5.69,16.9C4.81,15.77 4.24,14.42 4.06,13H2.06M7.1,18.37L5.67,19.74C7.18,21 9.04,21.79 11,22V20C9.58,19.82 8.23,19.25 7.1,18.37M12.5,7V12.25L17,14.92L16.25,16.15L11,13V7H12.5Z"/></svg>
            </div>
            <div class="feature-title">Batch Processing</div>
            <div class="feature-desc">Process multiple files</div>
        </div>
        <div class="feature-card">
            <div class="feature-icon">
                <svg viewBox="0 0 24 24" fill="white"><path d="M5,20H19V18H5M19,9H15V3H9V9H5L12,16L19,9Z"/></svg>
            </div>
            <div class="feature-title">ZIP Download</div>
            <div class="feature-desc">Download all at once</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize bulk rebrander
    if 'unified_rebrander' not in st.session_state:
        st.session_state.unified_rebrander = UnifiedBulkRebrander()
    
    if 'processing_results' not in st.session_state:
        st.session_state.processing_results = None
    
    # Sidebar
    with st.sidebar:
        # System Status Section
        st.markdown("""
        <div class="sidebar-section">
            <div class="sidebar-section-title">System Status</div>
        """, unsafe_allow_html=True)
        
        if DOCX_AVAILABLE:
            st.markdown("""
            <div class="sidebar-status ready">
                <span class="sidebar-status-dot green"></span>
                Word Rebranding Ready
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="sidebar-status unavailable">
                <span class="sidebar-status-dot red"></span>
                Word Not Available
            </div>
            """, unsafe_allow_html=True)
        
        if PDF_ALL_DEPS_AVAILABLE:
            st.markdown("""
            <div class="sidebar-status ready">
                <span class="sidebar-status-dot green"></span>
                PDF Rebranding Ready
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="sidebar-status unavailable">
                <span class="sidebar-status-dot red"></span>
                PDF Not Available
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Supported Formats Section
        st.markdown("""
        <div class="sidebar-section">
            <div class="sidebar-section-title">Supported Formats</div>
            <div class="sidebar-list-item">Word Documents (.docx)</div>
            <div class="sidebar-list-item">PDF Documents (.pdf)</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Rebranding Process Section
        st.markdown("""
        <div class="sidebar-section">
            <div class="sidebar-section-title">Scope</div>
            <div style="font-size: 0.75rem; font-weight: 600; color: #6C757D; margin-bottom: 0.4rem; text-transform: uppercase;">Preserved</div>
            <div class="sidebar-list-item check">All content and text</div>
            <div class="sidebar-list-item check">Tables and data</div>
            <div class="sidebar-list-item check">Fonts and formatting</div>
            <div class="sidebar-list-item check">Charts and images</div>
            <div style="font-size: 0.75rem; font-weight: 600; color: #6C757D; margin: 0.75rem 0 0.4rem 0; text-transform: uppercase;">Updated</div>
            <div class="sidebar-list-item update">DXC branding elements</div>
            <div class="sidebar-list-item update">Logo placement</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Templates Section
        st.markdown("""
        <div class="sidebar-section">
            <div class="sidebar-section-title">Template</div>
            <div class="sidebar-template">
                <span class="sidebar-template-label">Word</span>
                <span class="sidebar-template-value">DXC Word A4 Dec 2025</span>
            </div>
            <div class="sidebar-template">
                <span class="sidebar-template-label">PDF</span>
                <span class="sidebar-template-value">DXC Dec 2025</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    # Main content
    st.markdown('<div class="section-header">Upload Documents</div>', unsafe_allow_html=True)
    
    uploaded_files = st.file_uploader(
        "Select Word (.docx) and PDF (.pdf) files",
        type=['docx', 'pdf'],
        accept_multiple_files=True,
        help="You can upload multiple Word and PDF documents at once. All will be rebranded with DXC Dec 2025 template.",
        label_visibility="collapsed"
    )
    
    if uploaded_files:
        # Show uploaded files summary
        st.markdown(f"""
        <div class="upload-success">
            <strong>{len(uploaded_files)}</strong> file(s) uploaded and ready for rebranding
        </div>
        """, unsafe_allow_html=True)
        
        # File list
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown("**Uploaded Files:**")
            for f in uploaded_files:
                file_ext = Path(f.name).suffix.lower()
                badge_class = "docx" if file_ext == ".docx" else "pdf"
                badge_text = "DOC" if file_ext == ".docx" else "PDF"
                st.markdown(f"""
                <div class="file-item">
                    <span class="file-type-badge {badge_class}">{badge_text}</span>
                    {f.name}
                </div>
                """, unsafe_allow_html=True)
        
        with col2:
            # Count
            docx_count = sum(1 for f in uploaded_files if Path(f.name).suffix.lower() == '.docx')
            pdf_count = sum(1 for f in uploaded_files if Path(f.name).suffix.lower() == '.pdf')
            
            st.markdown("**Summary:**")
            st.markdown(f"Word Documents: **{docx_count}**")
            st.markdown(f"PDF Documents: **{pdf_count}**")
        
        st.markdown("---")
        
        # Color selection option with visual color boxes
        st.markdown("**Shape/Box Color Theme Option:**")
        
        # Define color options with hex codes
        color_options = {
            'Orange': '#FF7E51',
            'Black': '#000000',
            'Light Blue': '#87CEEB',
            'Dark Blue': '#00008B'
        }
        
        # Display color boxes above radio buttons for visual reference
        cols = st.columns(4)
        for idx, (color_name, hex_code) in enumerate(color_options.items()):
            with cols[idx]:
                st.markdown(f"""
                <div style="text-align: center; padding: 8px; margin-bottom: 5px;">
                    <div style="width: 50px; height: 50px; background-color: {hex_code}; border: 2px solid #333; border-radius: 4px; margin: 0 auto 5px;"></div>
                    <div style="font-size: 12px; font-weight: 500;">{color_name}</div>
                    <div style="font-size: 10px; color: #666; margin-top: 2px;">{hex_code}</div>
                </div>
                """, unsafe_allow_html=True)
        
        # Radio buttons for selection
        replacement_color = st.radio(
            "Choose the color for converted purple shapes/boxes:",
            options=list(color_options.keys()),
            index=0,  # Default to Orange
            horizontal=True,
            help="Select whether purple shapes/boxes should be converted to Orange (default), Black, Light Blue, or Dark Blue. Text color will automatically adapt - white on dark backgrounds, black on light backgrounds."
        )
        replacement_color_lower = replacement_color.lower()
        
        st.caption("💡 **Tip:** Text color will automatically adapt - white text on dark backgrounds, black text on light backgrounds")
        
        st.markdown("---")
        
        # Process button
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            process_btn = st.button(
                "Rebrand All Documents",
                use_container_width=True,
                type="primary"
            )
        
        if process_btn:
            # Progress tracking
            progress_container = st.empty()
            progress_bar = st.progress(0)
            current_file_container = st.empty()
            
            def update_progress(message, progress, file_name):
                progress_container.info(message)
                progress_bar.progress(progress / 100)
                if file_name:
                    current_file_container.caption(f"Processing: {file_name}")
            
            st.session_state.unified_rebrander.set_progress_callback(update_progress)
            
            # Process files
            with st.spinner("Processing documents..."):
                results = st.session_state.unified_rebrander.process_files(uploaded_files, replacement_color_lower)
                st.session_state.processing_results = results
            
            # Clear progress
            progress_container.empty()
            progress_bar.empty()
            current_file_container.empty()
            
            st.rerun()
    
    # Display results
    if st.session_state.processing_results:
        results = st.session_state.processing_results
        
        st.markdown("---")
        st.markdown('<div class="section-header">Processing Results</div>', unsafe_allow_html=True)
        
        # Statistics
        stats = results['stats']
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total Files", stats['total'])
        with col2:
            st.metric("Successful", stats['success_count'])
        with col3:
            st.metric("Failed", stats['fail_count'])
        with col4:
            st.metric("Time (sec)", stats['processing_time'])
        
        # Success rate
        if stats['total'] > 0:
            success_rate = (stats['success_count'] / stats['total']) * 100
            st.progress(success_rate / 100)
            st.caption(f"Success Rate: {success_rate:.1f}%")
        
        # Successful files
        if results['successful']:
            st.markdown("---")
            st.markdown('<div class="section-header">Successfully Rebranded</div>', unsafe_allow_html=True)
            
            # Download all as ZIP
            if len(results['successful']) > 1:
                zip_data = st.session_state.unified_rebrander.create_zip_download(results['successful'])
                zip_filename = f"DXC_Rebranded_Documents_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
                
                st.download_button(
                    label=f"Download All ({len(results['successful'])} files) as ZIP",
                    data=zip_data,
                    file_name=zip_filename,
                    mime="application/zip",
                    use_container_width=True,
                    type="primary"
                )
                
                st.markdown("---")
                st.markdown("**Or download individually:**")
            
            # Individual downloads
            for original_name, file_bytes, new_name in results['successful']:
                file_ext = Path(new_name).suffix.lower()
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if file_ext == '.docx' else "application/pdf"
                badge_class = "docx" if file_ext == ".docx" else "pdf"
                badge_text = "DOC" if file_ext == ".docx" else "PDF"
                
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"""
                    <div class="file-item">
                        <span class="file-type-badge {badge_class}">{badge_text}</span>
                        <span><strong>{original_name}</strong> &rarr; {new_name}</span>
                    </div>
                    """, unsafe_allow_html=True)
                with col2:
                    st.download_button(
                        label="Download",
                        data=file_bytes,
                        file_name=new_name,
                        mime=mime,
                        key=f"download_{new_name}"
                    )
        
        # Failed files
        if results['failed']:
            st.markdown("---")
            st.markdown('<div class="section-header">Failed to Rebrand</div>', unsafe_allow_html=True)
            
            for file_name, error in results['failed']:
                st.markdown(f"""
                <div class="upload-error">
                    <strong>{file_name}</strong><br>
                    {error}
                </div>
                """, unsafe_allow_html=True)
        
        # Clear results button
        st.markdown("---")
        if st.button("Clear Results and Start New Batch"):
            st.session_state.processing_results = None
            st.rerun()
    
    # Footer
    st.markdown(f"""
    <div class="enterprise-footer">
        <div class="footer-top">
            <div class="footer-credits">
                <span>Developed by <strong>Mayank Kumar</strong></span>
                <span class="footer-separator">|</span>
                <span>Product Owner <strong>Kirti Var Bisht</strong> & <strong>Jasveen Kaur</strong></span>
                <span class="footer-separator">|</span>
                <span>Executive Sponsor <strong>Mark Parsonage-Kear</strong> </span>
                <span class="footer-separator">|</span>
                <span><strong>Cross Functional Capabilities - AI and Automation</strong></span>
            </div>
        </div>
        <div class="footer-bottom">
            <div class="footer-copyright">DXC Technology &copy; {datetime.now().year}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
